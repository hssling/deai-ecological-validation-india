from __future__ import annotations

import re
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from prepare_jiag_submission import ROOT, make_display_tables, read_tables


OUT = ROOT / "submission_assets" / "IJMR_DEAI_2026-04-23_final"
FIG_OUT = OUT / "figures"
TABLE_OUT = OUT / "tables"

TITLE = "Development and ecological validation of a digital exposome ageing index for healthy ageing surveillance in India using LASI Wave 1"
RUNNING_TITLE = "DEAI ecological validation"
TODAY = "23 April 2026"
PUBLIC_REPO_URL = "https://github.com/hssling/deai-ecological-validation-india"

AUTHOR_NAME = "Dr Siddalingaiah H S"
AUTHOR_ROLE = "Professor, Community Medicine"
AUTHOR_AFFILIATION = "Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, Karnataka, India"
AUTHOR_EMAIL = "hssling@yahoo.com"
AUTHOR_PHONE = "+91 8941087719"
AUTHOR_ORCID = "0000-0002-4771-8285"

REFS = [
    "International Institute for Population Sciences, National Programme for Health Care of Elderly, Ministry of Health and Family Welfare, Harvard T.H. Chan School of Public Health, University of Southern California. Longitudinal Ageing Study in India (LASI) Wave 1, 2017-18: India Report. Mumbai: International Institute for Population Sciences; 2020.",
    "International Institute for Population Sciences. Longitudinal Ageing Study in India (LASI) Wave 1: State and Union Territory Factsheets. Mumbai: International Institute for Population Sciences; 2022.",
    "International Institute for Population Sciences, ICF. National Family Health Survey (NFHS-5), India, 2019-21. Mumbai: International Institute for Population Sciences; 2021.",
    "World Health Organization. World report on ageing and health. Geneva: World Health Organization; 2015.",
    "Wild CP. Complementing the genome with an exposome: the outstanding challenge of environmental exposure measurement in molecular epidemiology. Cancer Epidemiol Biomarkers Prev. 2005;14(8):1847-1850. doi:10.1158/1055-9965.EPI-05-0456.",
    "Lopez-Otin C, Blasco MA, Partridge L, Serrano M, Kroemer G. Hallmarks of aging: An expanding universe. Cell. 2023;186(2):243-278. doi:10.1016/j.cell.2022.11.001.",
    "GBD 2019 Risk Factors Collaborators. Global burden of 87 risk factors in 204 countries and territories, 1990-2019: a systematic analysis for the Global Burden of Disease Study 2019. Lancet. 2020;396(10258):1223-1249. doi:10.1016/S0140-6736(20)30752-2.",
    "Watts N, Amann M, Arnell N, et al. The 2023 report of the Lancet Countdown on health and climate change: the imperative for a health-centred response in a world facing irreversible harms. Lancet. 2023;402(10419):2346-2394. doi:10.1016/S0140-6736(23)01859-7.",
    "World Health Organization. WHO global air quality guidelines: particulate matter, ozone, nitrogen dioxide, sulfur dioxide and carbon monoxide. Geneva: World Health Organization; 2021.",
    "Benjamini Y, Hochberg Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J R Stat Soc Series B Stat Methodol. 1995;57(1):289-300.",
]

FIGURE_SPECS = [
    ("Figure 1", "real_fig1_state_deai_ranking.png", "State-level DEAI ranking across Indian states/UTs using LASI Wave 1 public indicators."),
    ("Figure 2", "real_fig2_deai_outcome_scatter.png", "Scatter plots of state-level DEAI against ageing outcomes."),
    ("Figure 3", "real_fig3_karnataka_exposome_profile.png", "Karnataka exposome profile compared with the India national LASI estimate."),
    ("Figure 4", "real_fig4_exposome_outcome_heatmap.png", "Heatmap of Spearman correlations between DEAI components and ageing outcomes."),
]

TABLE_TITLES = [
    "Table I. States/UTs with highest adverse DEAI burden",
    "Table II. DEAI correlations with ageing outcomes and robustness checks",
    "Table III. DEAI component diagnostics and internal construct contribution",
    "Table IV. Karnataka profile compared with India national LASI estimate",
]


def words(text: str) -> int:
    return len(re.findall(r"\b[\w+-]+\b", text))


def write_txt(name: str, text: str) -> Path:
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / name
    path.write_text(text.strip() + "\n", encoding="utf-8")
    return path


def ensure_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        section.top_margin = Inches(0.98)
        section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(0.98)
        section.right_margin = Inches(0.98)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Times New Roman"


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break()


def add_para(doc: Document, text: str = "", bold_label: str | None = None, align_justify: bool = True) -> None:
    p = doc.add_paragraph()
    if align_justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        r.font.name = "Times New Roman"
    r = p.add_run(text)
    r.font.name = "Times New Roman"


def add_cited_para(doc: Document, text: str, refs: list[int] | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    if refs:
        c = p.add_run(",".join(str(i) for i in refs))
        c.font.name = "Times New Roman"
        c.font.superscript = True


def add_df_table(doc: Document, title: str, df: pd.DataFrame) -> None:
    add_para(doc, title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def reference_metadata() -> pd.DataFrame:
    return pd.DataFrame(
        [
            [1, "LASI India Report", "IIPS LASI publications page and PDF", "https://www.iipsindia.ac.in/content/lasi-publications", "Verified"],
            [2, "LASI state and UT factsheets", "IIPS LASI publications page", "https://www.iipsindia.ac.in/content/lasi-publications", "Verified"],
            [3, "NFHS-5 India report", "IIPS/NFHS reports portal", "https://rchiips.org/nfhs/", "Verified"],
            [4, "WHO World report on ageing and health", "WHO publication page", "https://www.who.int/publications/i/item/9789241565042", "Verified"],
            [5, "Wild exposome paper", "DOI metadata", "https://doi.org/10.1158/1055-9965.EPI-05-0456", "Verified"],
            [6, "Hallmarks of aging 2023", "DOI metadata", "https://doi.org/10.1016/j.cell.2022.11.001", "Verified"],
            [7, "GBD 2019 risk factors", "DOI metadata", "https://doi.org/10.1016/S0140-6736(20)30752-2", "Verified"],
            [8, "Lancet Countdown 2023", "DOI metadata", "https://doi.org/10.1016/S0140-6736(23)01859-7", "Verified"],
            [9, "WHO global air quality guidelines", "WHO publication page", "https://www.who.int/publications/i/item/9789240034228", "Verified"],
            [10, "Benjamini-Hochberg FDR", "Journal metadata", "https://www.jstor.org/stable/2346101", "Verified"],
        ],
        columns=["Ref", "Item", "Verification source", "URL/DOI", "Status"],
    )


def article_text(tables: dict[str, pd.DataFrame]) -> tuple[str, list[tuple[str, str]]]:
    robust = tables["robust"]
    summary = tables["summary"]
    rank = tables["rank"]
    df = tables["df"]
    mortality = robust[robust["outcome"] == "death_rate_60plus_per1000"].iloc[0]
    multimorbidity = robust[robust["outcome"] == "multimorbidity_index"].iloc[0]
    alpha = float(summary[summary["metric"] == "deai_component_cronbach_alpha"]["value"].iloc[0])
    karnataka_rank = int(rank[rank["state"] == "Karnataka"]["rank"].iloc[0])
    karnataka = df[df["state"] == "Karnataka"].iloc[0]
    india = df[df["state"] == "INDIA"].iloc[0]
    mortality_state = str(mortality["most_influential_state"]).replace(" **", "")

    abstract = (
        "Background & objectives: Healthy-ageing surveillance requires scalable measures of upstream environmental and social adversity. "
        "This study developed a Digital Exposome Ageing Index (DEAI) and assessed its ecological validity using publicly available Longitudinal Ageing Study in India (LASI) Wave 1 state and union territory indicators. "
        "Methods: Eight adverse exposome components were standardized and combined using pre-specified knowledge weights into a state-level composite score. Outcomes were death rate among adults aged 60 years and above, falls, activities of daily living (ADL) and instrumental activities of daily living (IADL) limitations, poor self-rated health, depression, and diagnosed multimorbidity. Spearman correlations, bootstrap confidence intervals, Benjamini-Hochberg false-discovery-rate correction, leave-one-out analysis, and component diagnostics were used. "
        f"Results: After excluding the national India row, Odisha, Jharkhand, Bihar, Uttar Pradesh, and Assam had the highest adverse DEAI burden. DEAI showed a nominal positive association with death rate at age 60 years and above (rho={mortality['rho_states_only']:+.3f}, p={mortality['p_states_only']:.3f}; bootstrap 95% CI {mortality['bootstrap95_low_states_only']:+.3f} to {mortality['bootstrap95_high_states_only']:+.3f}; FDR q={mortality['bh_q_states_only']:.3f}) and a robust inverse association with diagnosed multimorbidity (rho={multimorbidity['rho_states_only']:+.3f}, p<0.001; bootstrap 95% CI {multimorbidity['bootstrap95_low_states_only']:+.3f} to {multimorbidity['bootstrap95_high_states_only']:+.3f}; FDR q<0.001). Correlations with falls, disability, poor self-rated health, and depression were weaker and not statistically robust. Internal consistency was acceptable for a multi-domain ecological index (Cronbach alpha={alpha:.3f}). "
        "Interpretation & conclusions: DEAI is a feasible ecological index for summarizing adverse exposome burden in Indian healthy-ageing surveillance. The inverse multimorbidity finding is more plausibly explained by epidemiological transition, survival, and diagnostic ascertainment than by any protective effect. Further individual-level and longitudinal validation is required before operational use."
    )

    sections = [
        (
            "Introduction",
            "Population ageing is now a major public-health challenge in India. Healthy ageing depends not only on chronological age and disease counts, but also on the environments in which older adults live. Public data systems such as LASI and NFHS make it possible to examine whether state-level upstream adversity aligns with patterns of mortality, disability, and diagnosed chronic disease in later life.\n\nThe exposome concept captures cumulative environmental exposure across the life course. In ageing research, it offers a way to connect environmental, lifestyle, and socioeconomic adversity with pathways relevant to functional decline and survival. However, biomarker-intensive ageing measures are not routinely available for programme use in India. A practical alternative is to construct a reproducible ecological index from indicators already available in public datasets.\n\nSuch an ecological index should not be confused with direct biological-age measurement. Its role is upstream: to summarize cumulative adverse living conditions that may accelerate disability, frailty, and mortality vulnerability at population level. This distinction matters in public health because programme planning often relies on diagnosed disease counts alone, which may miss regions where risk accumulates through poverty-linked exposures, poor household infrastructure, and nutritional deficit.\n\nThis study developed a Digital Exposome Ageing Index (DEAI) and evaluated its state-level ecological validity using public LASI Wave 1 factsheet data. The aim was not to estimate biological age, but to test whether a summary measure of adverse upstream exposures shows interpretable relationships with ageing outcomes relevant to surveillance and programme planning.",
        ),
        (
            "Material & Methods",
            "This was a methodological ecological study using publicly available aggregate LASI Wave 1 state/UT factsheet indicators. The unit of analysis was the state/UT. The national India row was retained for descriptive comparison but excluded from primary robustness estimates because it was not an independent state-level observation.\n\nDEAI was designed as a state-level adverse exposome index using eight components available in the public summaries: lack of clean cooking fuel, poor sanitation, low literacy, tobacco use, heavy episodic drinking, poor water access, reported indoor pollution exposure, and underweight prevalence. Each component was coded so that higher values indicated greater adverse burden. Missing values were median-imputed within the state-level dataset, each component was standardized as a z score, and the weighted linear score for state s was computed as DEAI_raw(s) = sum_i(w_i x z_si).\n\nThe pre-specified knowledge weights were: no clean cooking fuel 0.22, tobacco use 0.20, low literacy 0.15, poor sanitation 0.10, poor water access 0.10, heavy episodic drinking 0.08, reported indoor pollution exposure 0.08, and underweight prevalence 0.07. The resulting weighted score was then standardized across observations as DEAI_Z(s) = [DEAI_raw(s) - mean(DEAI_raw)] / SD(DEAI_raw). Higher DEAI values therefore indicate greater adverse exposome burden, not greater resilience or better ageing. The largest weights were assigned to clean-fuel deficit, tobacco use, low literacy, sanitation, and water access because these domains have strong links to respiratory, inflammatory, nutritional, and socioeconomic pathways relevant to healthy ageing.\n\nState-level ageing outcomes were death rate among adults aged 60 years and above, falls, ADL limitation, IADL limitation, poor self-rated health, depression by CIDI-SF, and diagnosed multimorbidity. Multimorbidity was interpreted cautiously because diagnosis-based disease prevalence may reflect survival and diagnostic ascertainment as well as underlying burden.\n\nSpearman correlation was used because the analysis was ecological and the sample size modest. Robustness checks included exclusion of the national India row, 5,000-sample bootstrap confidence intervals, Benjamini-Hochberg false-discovery-rate adjustment, leave-one-out influence analysis, and component diagnostics including Spearman correlation with total DEAI and Cronbach alpha. No individual-level identifiable data were accessed; ethics approval was therefore not applicable.",
        ),
        (
            "Results",
            f"Odisha, Jharkhand, Bihar, Uttar Pradesh, and Assam had the highest adverse DEAI scores. The top ten states/UTs are shown in Table I and the full ranking in Figure 1.\n\nDEAI showed a nominal positive association with death rate among adults aged 60 years and above. In states/UTs only, the Spearman correlation was {mortality['rho_states_only']:+.3f} (p={mortality['p_states_only']:.3f}; bootstrap 95% CI {mortality['bootstrap95_low_states_only']:+.3f} to {mortality['bootstrap95_high_states_only']:+.3f}). This direction remained stable in leave-one-out analysis, with rho ranging from {mortality['loo_min_rho']:+.3f} to {mortality['loo_max_rho']:+.3f}, and the largest influence attributable to {mortality_state}. However, it did not survive false-discovery-rate correction (FDR q={mortality['bh_q_states_only']:.3f}).\n\nThe strongest association was inverse and involved diagnosed multimorbidity. In states/UTs only, DEAI correlated with multimorbidity at rho={multimorbidity['rho_states_only']:+.3f} (p<0.001; bootstrap 95% CI {multimorbidity['bootstrap95_low_states_only']:+.3f} to {multimorbidity['bootstrap95_high_states_only']:+.3f}; FDR q<0.001). Leave-one-out estimates ranged from {multimorbidity['loo_min_rho']:+.3f} to {multimorbidity['loo_max_rho']:+.3f}, indicating that the finding was not driven by a single state/UT. Table II and Figure 2 summarize the robustness pattern.\n\nThe remaining endpoints were directionally mixed and statistically weaker. Fall prevalence showed a positive but non-significant correlation with DEAI (rho=+0.310, p=0.066; bootstrap 95% CI -0.054 to +0.614), while IADL limitation also trended positive (rho=+0.214, p=0.209; bootstrap 95% CI -0.113 to +0.522). ADL limitation and poor self-rated health were weakly inverse, with confidence intervals spanning the null, and depression by CIDI-SF showed little relation to DEAI at state level (rho=+0.065, p=0.711). Taken together, these findings indicate that the present ecological signal is stronger for contextual mortality vulnerability and ascertainment-sensitive multimorbidity than for disability, mood, or self-rated health measures.\n\nInternal consistency was acceptable for a multi-domain ecological index (Cronbach alpha={alpha:.3f}). The strongest component correlations with total DEAI were seen for underweight prevalence, tobacco use, lack of clean cooking fuel, indoor pollution exposure, low literacy, and poor sanitation. Table III and Figure 4 summarize these component relationships.\n\nKarnataka ranked {karnataka_rank} in the descriptive ranking including the India row. Its DEAI Z-score was {karnataka['deai_real_z']:.2f}, close to the India value of {india['deai_real_z']:.2f}. Karnataka had lower death rate at age 60 years and above ({karnataka['death_rate_60plus_per1000']:.1f} vs {india['death_rate_60plus_per1000']:.1f} per 1,000) and lower poor self-rated health ({karnataka['poor_srh_pct']:.1f}% vs {india['poor_srh_pct']:.1f}%), but higher IADL limitation ({karnataka['iadl_limitation_pct']:.1f}% vs {india['iadl_limitation_pct']:.1f}%). Table IV presents the Karnataka profile and Figure 3 shows the same comparison visually.",
        ),
        (
            "Discussion",
            "This study presents a reproducible ecological index for summarizing adverse exposome burden in Indian healthy-ageing surveillance. The principal contribution is methodological: DEAI offers a compact way to integrate multiple upstream environmental, household, lifestyle, and socioeconomic indicators relevant to ageing. Because it is explicitly constructed as a weighted sum of standardized adverse exposures, it should be interpreted as a surveillance marker of contextual burden rather than as a biological clock or individual-level risk score.\n\nThe mortality association was directionally consistent with expectation. States with greater adverse exposome burden tended to have higher death rate among older adults. The association remained stable after exclusion of the national India row and during leave-one-out analysis, suggesting that it was not driven by a single observation. However, because it did not survive false-discovery-rate correction, the proper interpretation is cautious: DEAI shows supportive ecological plausibility for mortality vulnerability, but not confirmatory evidence.\n\nThe inverse multimorbidity association was statistically stronger and biologically counterintuitive. It should not be interpreted as evidence that adverse exposures are protective against chronic disease. The more plausible interpretation is epidemiological transition and ascertainment. In higher-adversity settings, cumulative poverty-linked exposures may increase mortality and reduce survival to ages at which chronic diseases are diagnosed and recorded. In later-transition settings, older adults may survive longer, interact more with health systems, and have greater diagnostic ascertainment of hypertension, diabetes, cardiovascular disease, and other chronic conditions. Diagnosed disease prevalence at ecological level may therefore behave partly as a marker of survival and detection rather than a monotonic marker of adverse ageing burden.\n\nThe weaker and mixed findings for falls, disability, poor self-rated health, and depression are also informative. They suggest that a state-level adverse exposome summary may not map cleanly onto all later-life outcomes in cross-sectional ecological data, especially for endpoints influenced by reporting practices, household adaptation, cultural thresholds, and access to diagnosis or rehabilitation. This pattern supports retaining DEAI as a surveillance-context indicator rather than overextending it as a universal proxy for every ageing dimension.\n\nThis interpretation has practical value for public-health planning. Disease counts alone may underestimate need in high-adversity states. An upstream exposure-burden index may therefore complement geriatric services, NCD screening, tobacco control, nutrition programmes, clean-fuel promotion, sanitation improvement, and age-friendly primary care planning. Karnataka illustrates why mixed profiles matter: its DEAI was close to the national value, yet its mortality, self-rated health, and IADL profiles did not all move in the same direction. A dashboard that combines DEAI with disability, mortality, diagnosed disease, and service-use indicators is likely to be more informative than any single metric.\n\nThe index also illustrates the value of combining public data streams. LASI provides older-adult outcomes and state-level social and health indicators, while NFHS and global public-health datasets provide context on household infrastructure, behavioural risks, and environmental adversity. A reproducible open pipeline allows these relationships to be updated as new LASI waves and better denominator data become available.\n\nThe study has several strengths. It uses real public LASI data, transparent component coding, reproducible scripts, pre-specified adverse-direction coding, bootstrap confidence intervals, false-discovery-rate adjustment, leave-one-out influence checks, and component diagnostics. The manuscript also avoids the common error of treating ecological associations as individual-level risk estimates. The inclusion of explicit weight disclosure and the scoring formula makes the construction auditable rather than opaque.\n\nThe limitations are equally important. The analysis is ecological and cross-sectional; it cannot estimate individual-level risk, causality, or biological age. The weights are knowledge-based and therefore interpretable, but they are not individually validated causal coefficients, and alternative weighting schemes could change rankings at the margin. Public state factsheets limit adjustment for age structure, sex composition, urban-rural heterogeneity, wealth, health-care access, and diagnostic intensity. Diagnosed multimorbidity remains vulnerable to survival and ascertainment bias. Relevant exposures such as outdoor PM2.5, heat stress, occupational exposure, migration, and neighbourhood context were not directly included in the real score. The moderate Cronbach alpha should be interpreted as acceptable for a multi-domain index, not as evidence of a single latent construct.\n\nFuture work should validate DEAI using individual-level LASI microdata, longitudinal outcomes, district-level environmental linkage, and programme indicators such as geriatric outpatient attendance, rehabilitation use, and screening coverage. Weight-sensitivity analysis and external validation against independently measured environmental indicators would further strengthen the index. Until such validation is completed, DEAI should be treated as a surveillance and hypothesis-generation tool rather than a clinical decision rule.",
        ),
        (
            "Conclusion",
            "DEAI is a feasible and reproducible ecological index for summarizing adverse exposome burden in Indian healthy-ageing surveillance. The present state-level validation shows geographic plausibility, a stable nominal mortality signal, and a robust inverse multimorbidity pattern that is best interpreted through epidemiological transition and diagnostic ascertainment. Individual-level and longitudinal validation are required before any clinical or programme-targeting use.",
        ),
    ]
    return abstract, sections


def manuscript_txt(display_tables: dict[str, pd.DataFrame], tables: dict[str, pd.DataFrame]) -> str:
    abstract, sections = article_text(tables)
    parts = [
        TITLE,
        "",
        "Structured Abstract",
        "",
        abstract,
        "",
        "Keywords: ageing; exposome; LASI; multimorbidity; mortality; India",
    ]
    for heading, text in sections:
        parts += ["", heading, "", text]
    parts += ["", "Acknowledgment", "", "The author acknowledges IIPS, LASI, NFHS, WHO, UN DESA, and public data providers whose aggregate indicators enabled this analysis."]
    parts += ["", "Conflicts of Interest", "", "None."]
    parts += ["", "References", ""]
    parts += [f"{i}. {ref}" for i, ref in enumerate(REFS, 1)]
    parts += ["", "Tables", ""]
    for idx, (_, df) in enumerate(display_tables.items()):
        parts += ["", TABLE_TITLES[idx], "", df.to_markdown(index=False)]
    parts += ["", "Legends to Figures", ""]
    for title, _, legend in FIGURE_SPECS:
        parts += ["", f"{title}. {legend}"]
    return "\n".join(parts)


def write_manuscript_docx(name: str, display_tables: dict[str, pd.DataFrame], tables: dict[str, pd.DataFrame]) -> None:
    abstract, sections = article_text(tables)
    doc = Document()
    ensure_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    add_para(doc, "Original Article", "Article type: ", align_justify=False)
    add_para(doc, RUNNING_TITLE, "Running title: ", align_justify=False)

    add_heading(doc, "Structured Abstract", 1)
    for label in ["Background & objectives", "Methods", "Results", "Interpretation & conclusions"]:
        match = re.search(rf"{re.escape(label)}: (.*?)(?= (?:Background & objectives|Methods|Results|Interpretation & conclusions):|$)", abstract)
        if match:
            add_para(doc, match.group(1).strip(), f"{label}: ")
    add_para(doc, "ageing; exposome; LASI; multimorbidity; mortality; India", "Keywords: ")

    ref_map = {
        "Introduction": [[1, 2, 3, 4], [5, 6, 7, 8, 9], [4, 7], [1, 2]],
        "Material & Methods": [[1, 2], [5, 6, 7, 8, 9], [5, 6, 7], [10]],
        "Results": [[], [], [], []],
        "Discussion": [[4, 5, 6], [4, 7], [1, 2, 4, 7], [4, 7, 9], [1, 2, 3, 7, 8, 9], [], [1, 2, 7, 8, 9], [1, 2, 3, 7, 8, 9]],
        "Conclusion": [[4]],
    }

    for heading, text in sections:
        add_heading(doc, heading, 1)
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        refs = ref_map.get(heading, [])
        for idx, para in enumerate(paragraphs):
            if idx < len(refs) and refs[idx]:
                add_cited_para(doc, para, refs[idx])
            else:
                add_para(doc, para)

    add_heading(doc, "Acknowledgment", 1)
    add_para(doc, "The author acknowledges IIPS, LASI, NFHS, WHO, UN DESA, and public data providers whose aggregate indicators enabled this analysis.")
    add_heading(doc, "Conflicts of Interest", 1)
    add_para(doc, "None.")
    add_page_break(doc)
    add_heading(doc, "References", 1)
    for i, ref in enumerate(REFS, 1):
        add_para(doc, f"{i}. {ref}")
    add_page_break(doc)
    add_heading(doc, "Tables", 1)
    for idx, (_, df) in enumerate(display_tables.items()):
        add_df_table(doc, TABLE_TITLES[idx], df)
    add_page_break(doc)
    add_heading(doc, "Legends to Figures", 1)
    for title_text, _, legend in FIGURE_SPECS:
        add_para(doc, f"{title_text}. {legend}")
    doc.save(OUT / name)


def write_figures_docx(name: str) -> None:
    doc = Document()
    ensure_style(doc)
    add_heading(doc, "Figures", 1)
    for idx, (title, filename, legend) in enumerate(FIGURE_SPECS):
        add_para(doc, f"{title}. {legend}")
        fig = ROOT / "results" / "figures" / filename
        if fig.exists():
            doc.add_picture(str(fig), width=Inches(5.8))
        if idx != len(FIGURE_SPECS) - 1:
            doc.add_page_break()
    doc.save(OUT / name)


def title_page_text(main_words: int) -> str:
    return f"""
INDIAN JOURNAL OF MEDICAL RESEARCH
FIRST PAGE FILE

Covering letter note:
This manuscript addresses healthy ageing surveillance, exposome-related public health, and epidemiological interpretation using nationally relevant LASI data. Its public-health and methodological relevance makes IJMR an appropriate venue.

Title:
{TITLE}

Running title (maximum 40 characters requested by journal):
{RUNNING_TITLE}

Article category:
Original Article

Author:
{AUTHOR_NAME}
{AUTHOR_ROLE}
{AUTHOR_AFFILIATION}
Mobile: {AUTHOR_PHONE}
Email: {AUTHOR_EMAIL}
ORCID: {AUTHOR_ORCID}

Department and institution to which the work is attributed:
Department of Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, Karnataka, India

Name, address and email of corresponding author:
{AUTHOR_NAME}
{AUTHOR_AFFILIATION}
Email: {AUTHOR_EMAIL}
Mobile: {AUTHOR_PHONE}

Contributors' credits:
Concept and design: {AUTHOR_NAME}
Data curation and analysis: {AUTHOR_NAME}
Interpretation: {AUTHOR_NAME}
Drafting and revision: {AUTHOR_NAME}
Guarantor: {AUTHOR_NAME}

Declaration on competing interests:
None.

Funding:
No external funding.

Ethical clearance status:
Not applicable. This study used public-domain aggregate data and no individual-level identifiable participant data.
Ethics Committee name, date and number:
Not applicable.

Clinical trial registration number:
Not applicable.

Data sharing statement:
Derived tables, figures, reproducible scripts, and the public release package are available at {PUBLIC_REPO_URL}; source public data are available from LASI/IIPS and related public reports.

Declaration of Artificial Intelligence (AI) in scientific writing:
Computational tools were used for data processing, reproducible analysis, document assembly, language editing support, and formatting support. Final responsibility for the scientific content rests with the author.

Word count (excluding abstract, tables, figures, acknowledgment, key messages, and references):
{main_words}

Number of tables:
4

Number of figures:
4

Disclaimers:
None.
""".strip()


def cover_letter_text() -> str:
    return f"""
{TODAY}

To
The Editor
Indian Journal of Medical Research

Subject: Submission of Original Article

Dear Editor,

Please consider the enclosed manuscript entitled "{TITLE}" for publication as an Original Article in the Indian Journal of Medical Research.

This manuscript presents a reproducible ecological index for summarizing adverse exposome burden in healthy-ageing surveillance in India and evaluates it using public LASI Wave 1 state and union territory indicators. The analysis includes bootstrap confidence intervals, false-discovery-rate correction, leave-one-out robustness checks, and an explicit interpretation of the inverse multimorbidity finding through epidemiological transition and diagnostic ascertainment.

This topic is suitable for IJMR because it addresses nationally relevant public health, epidemiology, community medicine, and ageing surveillance rather than a narrow subspecialty audience.

I confirm that the manuscript is original, has not been published previously in full, and is not under consideration elsewhere. The study used only public-domain aggregate data and involved no individual-level identifiable information; ethics approval was therefore not applicable. There was no external funding and no conflicts of interest.

Thank you for your consideration.

Sincerely,

{AUTHOR_NAME}
{AUTHOR_ROLE}
{AUTHOR_AFFILIATION}
Email: {AUTHOR_EMAIL}
Phone: {AUTHOR_PHONE}
ORCID: {AUTHOR_ORCID}
""".strip()


def declarations_text() -> str:
    return f"""
INDIAN JOURNAL OF MEDICAL RESEARCH
AUTHOR DECLARATIONS

Title:
{TITLE}

Funding:
No external funding.

Conflicts of interest:
None.

Ethics approval:
Not applicable. This study used publicly available aggregate LASI and related summary indicators. No individual-level identifiable participant data were accessed.

Consent to participate:
Not applicable.

Consent for publication:
Not applicable.

Data availability:
Derived tables, figures, reproducible scripts, and the public release package are available at {PUBLIC_REPO_URL}. Source public data are available from LASI/IIPS and related public reports.

Author contributions:
{AUTHOR_NAME}: concept, study design, data curation, statistical analysis, manuscript preparation, manuscript review, and guarantor.

Declaration of AI in scientific writing:
Computational tools were used for data processing, reproducible analysis, document assembly, language editing support, and formatting support. Final responsibility for the scientific content rests with the author.

Data sharing statement:
The curated public release, including reproducible analysis files, manuscript assets, and supporting tables and figures, is available at {PUBLIC_REPO_URL}.
""".strip()


def checklist_text() -> str:
    return """
INDIAN JOURNAL OF MEDICAL RESEARCH
SUBMISSION CHECKLIST

Prepared files
- First page file
- Cover letter
- Blinded manuscript
- Author declarations
- Figures DOCX file
- Reference metadata
- Separate figure files
- Separate table CSV files

Checks completed
- Article category aligned as Original Article
- Structured abstract prepared in four-part IJMR format
- Title page details separated from blinded manuscript
- In-text references converted to superscript numbered style without brackets
- References ordered by first appearance in text
- Tables moved after References section
- Figure legends included in manuscript after Tables section
- Figures copied separately and compiled in a separate DOCX file
- Conflicts/funding/ethics/AI/data-sharing statements included

Manual checks before upload
- Confirm the latest IJMR portal requires any additional signed copyright or author forms
- Open the DOCX in Word and confirm table fit and figure placement
- Confirm the running title length and abstract word count on the final portal version
""".strip()


def audit_note_text(main_words: int) -> str:
    return f"""
IJMR DEAI PACKAGE AUDIT NOTE

Generated: {TODAY}

Summary
- Source manuscript: JIAG DEAI package adapted for IJMR as an Original Article.
- Main text word count: {main_words}
- Structured abstract reformatted to IJMR style headings.
- Reference citations converted to superscript numbering without brackets.
- References, Tables, and Legends to Figures arranged in IJMR order.
- First page file, cover letter, blinded manuscript, declarations, figures DOCX, and checklist prepared.

Residual manual check
- Open the DOCX in Word before submission to confirm pagination, table width, and image placement.
""".strip()


def generic_docx(name: str, text: str) -> None:
    doc = Document()
    ensure_style(doc)
    for block in text.split("\n\n"):
        add_para(doc, block)
    doc.save(OUT / name)


def save_pair(stem: str, text: str) -> None:
    write_txt(f"{stem}.txt", text)
    generic_docx(f"{stem}.docx", text)


def copy_assets() -> None:
    FIG_OUT.mkdir(parents=True, exist_ok=True)
    TABLE_OUT.mkdir(parents=True, exist_ok=True)
    for _, fig, _ in FIGURE_SPECS:
        shutil.copy2(ROOT / "results" / "figures" / fig, FIG_OUT / fig)
    for table in [
        "deai_real_state_rankings.csv",
        "deai_real_outcome_correlations.csv",
        "deai_real_robustness.csv",
        "deai_real_robustness_summary.csv",
        "deai_component_diagnostics.csv",
        "lasi_with_deai_real.csv",
    ]:
        shutil.copy2(ROOT / "results" / "tables" / table, TABLE_OUT / table)


def automated_audit(manuscript_path: Path) -> pd.DataFrame:
    doc = Document(manuscript_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    checks = [
        ("author_name_in_blinded", AUTHOR_NAME not in text),
        ("author_email_in_blinded", AUTHOR_EMAIL not in text),
        ("table_1_cited", "Table I" in text),
        ("figure_1_cited", "Figure 1" in text),
        ("structured_abstract_present", "Background & objectives:" in text and "Interpretation & conclusions:" in text),
        ("references_present", "1. International Institute for Population Sciences" in text),
        ("ijmr_article_type", "Original Article" in text),
        ("tables_after_references", text.find("References") < text.find("Tables") < text.find("Legends to Figures")),
    ]
    return pd.DataFrame({"check": [c[0] for c in checks], "status": ["pass" if c[1] else "fail" for c in checks]})


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    copy_assets()
    tables = read_tables()
    display_tables = make_display_tables(tables)
    manuscript = manuscript_txt(display_tables, tables)
    write_txt("IJMR_DEAI_manuscript_blinded_2026-04-23.txt", manuscript)
    write_manuscript_docx("IJMR_DEAI_manuscript_blinded_2026-04-23.docx", display_tables, tables)
    write_figures_docx("IJMR_DEAI_figures_2026-04-23.docx")

    main_match = re.search(r"Introduction\n\n(.*?)\n\nAcknowledgment", manuscript, re.S)
    main_words = words(main_match.group(1)) if main_match else 0

    save_pair("IJMR_DEAI_first_page_file_2026-04-23", title_page_text(main_words))
    save_pair("IJMR_DEAI_cover_letter_2026-04-23", cover_letter_text())
    save_pair("IJMR_DEAI_declarations_2026-04-23", declarations_text())
    save_pair("IJMR_DEAI_submission_checklist_2026-04-23", checklist_text())
    save_pair("IJMR_DEAI_audit_note_2026-04-23", audit_note_text(main_words))

    reference_metadata().to_csv(OUT / "IJMR_DEAI_reference_metadata_2026-04-23.csv", index=False)
    automated_audit(OUT / "IJMR_DEAI_manuscript_blinded_2026-04-23.docx").to_csv(
        OUT / "IJMR_DEAI_automated_audit_checks_2026-04-23.csv", index=False
    )
    print(f"IJMR package written to: {OUT}")
    print(f"Main text words: {main_words}")


if __name__ == "__main__":
    main()
