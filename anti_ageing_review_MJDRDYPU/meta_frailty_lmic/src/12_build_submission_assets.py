from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "results" / "tables"
DOCS = ROOT / "docs"
ASSET_DIR = ROOT / "submission_assets" / "IJMR_FRAILTY_INTERVENTIONS_AUDIT_READY_2026-05-18"
FIG_DIR = ASSET_DIR / "figures"

TITLE = (
    "Community-deliverable interventions for physical frailty in older adults: "
    "a systematic review, conditional component network meta-analysis, and "
    "implementation-readiness mapping for Indian primary care"
)
SHORT_TITLE = "Frailty interventions for Indian primary care"
AUTHOR_1 = "Dr Siddalingaiah H S"
AUTHOR_2 = "Dr Chandrakala D"


def load_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES / name).fillna("")


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run(value)


def add_table_from_df(doc: Document, df: pd.DataFrame, max_rows: int | None = None) -> None:
    if max_rows is not None:
        df = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])


def save_doc(doc: Document, name: str) -> None:
    doc.save(ASSET_DIR / name)


def make_prisma_progress_figure() -> Path:
    search_log = load_csv("search_log.csv").iloc[0]
    counts = {
        "PubMed hits reported": int(search_log["hits_reported"]),
        "Records downloaded": int(search_log["records_downloaded"]),
        "Title/abstract potential": int(
            load_csv("title_abs_screening_counts.csv")
            .query("screen_decision == 'include_title_abs'")["n"]
            .iloc[0]
        ),
        "Primary full-text queue": int(
            load_csv("fulltext_priority_counts.csv")
            .query("fulltext_priority == 'A_primary_fulltext'")["n"]
            .iloc[0]
        ),
        "High-confidence extraction": len(load_csv("high_confidence_extraction_queue.csv")),
        "PMC full texts mined": int(
            load_csv("fulltext_access_status_counts.csv")
            .query("fulltext_access_status == 'pmc_fulltext_available'")["n"]
            .iloc[0]
        ),
        "Final included studies": 0,
    }
    labels = list(counts.keys())
    values = list(counts.values())

    fig, ax = plt.subplots(figsize=(9, 6))
    ax.axis("off")
    y = 0.95
    for i, (label, value) in enumerate(zip(labels, values)):
        box = dict(boxstyle="round,pad=0.35", facecolor="#f2f5f7", edgecolor="#4b5563", linewidth=1)
        ax.text(0.5, y, f"{label}\n{value}", ha="center", va="center", fontsize=10, bbox=box)
        if i < len(labels) - 1:
            ax.annotate("", xy=(0.5, y - 0.085), xytext=(0.5, y - 0.035), arrowprops=dict(arrowstyle="->", lw=1.2))
        y -= 0.14
    ax.text(
        0.5,
        0.02,
        "Progress flow only: final full-text exclusions and included studies are pending author verification.",
        ha="center",
        va="bottom",
        fontsize=9,
        color="#7f1d1d",
    )
    out = FIG_DIR / "figure1_screening_progress_flow.png"
    fig.tight_layout()
    fig.savefig(out, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return out


def make_node_figure() -> Path:
    nodes = load_csv("network_prelim_node_counts.csv")
    nodes = nodes.sort_values("candidate_records", ascending=True)
    fig, ax = plt.subplots(figsize=(8, 4.8))
    ax.barh(nodes["preliminary_node"], nodes["candidate_records"], label="High-confidence candidates", color="#2563eb")
    ax.barh(nodes["preliminary_node"], nodes["pmc_mined_records"], label="PMC mined full texts", color="#f97316")
    ax.set_xlabel("Records")
    ax.set_title("Preliminary intervention-node distribution")
    ax.legend(loc="lower right")
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    out = FIG_DIR / "figure2_preliminary_node_distribution.png"
    fig.savefig(out, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return out


def build_first_page() -> None:
    doc = Document()
    set_normal_style(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph()
    add_kv(doc, "Article type", "Systematic review and meta-analysis; NMA conditional on final network feasibility")
    add_kv(doc, "Target journal", "Indian Journal of Medical Research")
    add_kv(doc, "Running title", SHORT_TITLE)
    add_kv(doc, "Authors", f"{AUTHOR_1}; {AUTHOR_2}")
    add_kv(doc, "Affiliations", "To be inserted and verified by authors before portal upload")
    add_kv(doc, "Corresponding author", f"{AUTHOR_1}; email, postal address and telephone pending author confirmation")
    add_kv(doc, "Word count", "Pending final manuscript")
    add_kv(doc, "Abstract word count", "Pending final manuscript")
    add_kv(doc, "Tables and figures", "Current audit package contains 2 figures and supplementary tables; final count pending synthesis")
    add_kv(doc, "Protocol registration", "PROSPERO/other review-registry number pending; required before IJMR submission")
    add_kv(doc, "Keywords", "frailty; older adults; exercise; nutrition; primary care; systematic review; network meta-analysis; India")
    doc.add_paragraph(
        "Submission-readiness note: this first page is structurally prepared but not portal-ready until registration, affiliations, "
        "correspondence details, final included studies, effect estimates and author declarations are confirmed."
    )
    save_doc(doc, "IJMR_frailty_first_page_2026-05-18.docx")


def build_declarations() -> None:
    doc = Document()
    set_normal_style(doc)
    add_heading(doc, "Declarations")
    add_kv(doc, "Manuscript title", TITLE)
    add_kv(doc, "Authors", f"{AUTHOR_1}; {AUTHOR_2}")
    add_heading(doc, "Author Contributions", 2)
    doc.add_paragraph(
        f"{AUTHOR_1}: Conceptualization, methodology, software-assisted literature workflow, formal analysis planning, "
        "data curation, writing - original draft, visualization, project administration, corresponding author role pending confirmation."
    )
    doc.add_paragraph(
        f"{AUTHOR_2}: Conceptualization, methodology, validation, investigation, writing - review and editing, supervision, "
        "second-reviewer verification role pending completion."
    )
    add_heading(doc, "Conflicts of Interest", 2)
    doc.add_paragraph("No conflicts of interest have been entered in the project files. Final author confirmation is required before submission.")
    add_heading(doc, "Funding", 2)
    doc.add_paragraph("No external funding has been entered in the project files. Final author confirmation is required before submission.")
    add_heading(doc, "Ethics Approval", 2)
    doc.add_paragraph(
        "This systematic review uses published aggregate data and does not require human-participant ethics approval. "
        "If local institutional policy requires exemption documentation, it should be added before submission."
    )
    add_heading(doc, "Data Availability", 2)
    doc.add_paragraph(
        "Search logs, screening outputs, extraction forms, audit reports and analysis scripts are maintained in the project workspace. "
        "Final de-identified extraction data and analysis code should be deposited or made available with the submission, subject to journal policy."
    )
    add_heading(doc, "Protocol Registration", 2)
    doc.add_paragraph("Protocol registration is pending. IJMR submission should not proceed until a registry number is available.")
    add_heading(doc, "Use of AI-Assisted Workflow", 2)
    doc.add_paragraph(
        "AI-assisted tools were used to structure search outputs, triage records, mine open-access full-text snippets and prepare audit-ready forms. "
        "Final study inclusion, exclusions, risk-of-bias judgements, numeric extraction, interpretation and submission approval require author verification."
    )
    save_doc(doc, "IJMR_frailty_declarations_2026-05-18.docx")


def build_figures_docx(fig1: Path, fig2: Path) -> None:
    doc = Document()
    set_normal_style(doc)
    add_heading(doc, "Figure File")
    doc.add_paragraph("Figure 1. Screening progress flow for the frailty intervention review.")
    doc.add_picture(str(fig1), width=Inches(6.5))
    doc.add_paragraph(
        "Caption: Records identified, screened and triaged through the current audit stage. This is not the final PRISMA flow diagram because full-text exclusions and final included studies remain pending."
    )
    doc.add_page_break()
    doc.add_paragraph("Figure 2. Preliminary intervention-node distribution.")
    doc.add_picture(str(fig2), width=Inches(6.5))
    doc.add_paragraph(
        "Caption: Distribution of high-confidence candidate records and PMC-mined full texts by preliminary intervention node. Nodes require confirmation after full-text extraction."
    )
    save_doc(doc, "IJMR_frailty_figures_2026-05-18.docx")


def build_supplementary() -> None:
    doc = Document()
    set_normal_style(doc)
    add_heading(doc, "Supplementary Appendix")
    doc.add_paragraph(TITLE)
    add_heading(doc, "Supplementary Table 1. PubMed Search Log", 2)
    add_table_from_df(doc, load_csv("search_log.csv"))
    add_heading(doc, "Supplementary Table 2. Title/Abstract Screening Counts", 2)
    add_table_from_df(doc, load_csv("title_abs_screening_counts.csv"))
    add_heading(doc, "Supplementary Table 3. Full-Text Priority Counts", 2)
    add_table_from_df(doc, load_csv("fulltext_priority_counts.csv"))
    add_heading(doc, "Supplementary Table 4. High-Confidence Triage Counts", 2)
    add_table_from_df(doc, load_csv("high_confidence_triage_counts.csv"))
    add_heading(doc, "Supplementary Table 5. Full-Text Access Status", 2)
    add_table_from_df(doc, load_csv("fulltext_access_status_counts.csv"))
    add_heading(doc, "Supplementary Table 6. Preliminary Network Feasibility", 2)
    add_table_from_df(doc, load_csv("network_feasibility_prelim.csv"))
    add_heading(doc, "Supplementary Table 7. Preliminary Node Counts", 2)
    add_table_from_df(doc, load_csv("network_prelim_node_counts.csv"))
    add_heading(doc, "Supplementary Methods Note", 2)
    doc.add_paragraph(
        "The current package documents a reproducible search, screening and extraction-preparation workflow. "
        "It does not report final pooled effects. Numeric outcome extraction, RoB 2 assessment, final component coding and network connectivity checks remain required."
    )
    save_doc(doc, "IJMR_frailty_supplementary_2026-05-18.docx")


def build_manuscript_scaffold() -> None:
    doc = Document()
    set_normal_style(doc)
    add_heading(doc, TITLE)
    add_heading(doc, "Abstract", 1)
    doc.add_paragraph(
        "Background: Physical frailty is a major barrier to healthy ageing and primary-care resilience in India. "
        "Objective: To synthesize randomized evidence on community-deliverable interventions for prefrail or frail older adults and map implementation readiness for Indian primary care. "
        "Methods: A PubMed search was executed on 2026-05-18. Final synthesis is pending protocol registration, full-text verification, numeric extraction and risk-of-bias assessment. "
        "Results: The current authenticated workflow identified 3425 PubMed hits, downloaded 3418 records, flagged 1129 title/abstract records as potentially relevant, and identified 179 high-confidence extraction candidates; 95 PMC full texts were mined as extraction aids. "
        "Conclusions: The evidence base appears feasible for full systematic review and probably pairwise meta-analysis. NMA remains conditional on final comparator connectivity and transitivity."
    )
    add_heading(doc, "Introduction", 1)
    doc.add_paragraph(
        "This scaffold is intentionally claim-limited. It should be converted into a full manuscript only after author-verified extraction and synthesis. "
        "The final report should follow PRISMA 2020 [1] and, if network meta-analysis proceeds, the PRISMA-NMA extension [2]."
    )
    add_heading(doc, "Methods", 1)
    doc.add_paragraph(
        "The target journal is Indian Journal of Medical Research, which considers systematic reviews including meta-analysis and requires authors to follow its technical submission process [3]. "
        "Protocol registration is pending and is required before IJMR submission."
    )
    doc.add_paragraph(
        "Randomized and cluster-randomized trials of community, home, outpatient or primary-care interventions for prefrail or frail older adults are the primary target. "
        "RoB 2 will be used for randomized trials after final outcome extraction [4]."
    )
    doc.add_paragraph("Figure 1 summarizes the current screening-progress flow. Figure 2 summarizes preliminary intervention nodes.")
    doc.add_paragraph("Table 1 summarizes authenticated workflow counts. Table 2 summarizes evidence gates before final synthesis.")
    add_heading(doc, "Results", 1)
    doc.add_paragraph("Final included studies, effect estimates, heterogeneity, certainty and NMA results are pending.")
    add_heading(doc, "Table 1. Authenticated Workflow Counts", 2)
    table1 = pd.DataFrame(
        [
            ["PubMed hits reported", "3425"],
            ["Records downloaded", "3418"],
            ["Title/abstract records flagged potentially relevant", "1129"],
            ["Primary full-text queue", "297"],
            ["High-confidence extraction candidates", "179"],
            ["PMC full texts mined", "95"],
            ["Publisher/library full texts required", "84"],
        ],
        columns=["Item", "Value"],
    )
    add_table_from_df(doc, table1)
    add_heading(doc, "Table 2. Evidence Gates Before Final Inference", 2)
    table2 = pd.DataFrame(
        [
            ["Protocol registration", "Pending", "Blocking before IJMR submission"],
            ["Final full-text inclusion", "Pending", "Required before final PRISMA"],
            ["Numeric extraction", "Pending", "Required before meta-analysis"],
            ["RoB 2", "Pending", "Required before interpretation"],
            ["NMA connectivity/transitivity", "Not assessed", "Required before NMA"],
        ],
        columns=["Gate", "Current status", "Implication"],
    )
    add_table_from_df(doc, table2)
    add_heading(doc, "Discussion", 1)
    doc.add_paragraph(
        "The main expected contribution is India-facing implementation readiness, not another broad global ranking of exercise and nutrition interventions. "
        "Any final conclusion must separate efficacy, safety, adherence and deliverability."
    )
    add_heading(doc, "Limitations of Current Scaffold", 1)
    doc.add_paragraph(
        "This document is not a completed manuscript. It contains authenticated workflow counts and submission structure, but lacks final full-text screening, effect sizes, RoB 2 judgements, certainty assessment and final reference metadata."
    )
    add_heading(doc, "References", 1)
    refs = [
        "Page MJ, McKenzie JE, Bossuyt PM, et al. The PRISMA 2020 statement: an updated guideline for reporting systematic reviews. BMJ. 2021;372:n71. doi:10.1136/bmj.n71.",
        "Hutton B, Salanti G, Caldwell DM, et al. The PRISMA extension statement for reporting of systematic reviews incorporating network meta-analyses of health care interventions. Ann Intern Med. 2015;162(11):777-784. doi:10.7326/M14-2385.",
        "Indian Journal of Medical Research. Instructions for authors. Available from: https://ijmr.org.in/for-authors/ Accessed May 18, 2026.",
        "Cochrane Methods. Risk of Bias 2 (RoB 2) tool. Available from: https://methods.cochrane.org/risk-bias-2 Accessed May 18, 2026.",
    ]
    for i, ref in enumerate(refs, start=1):
        doc.add_paragraph(f"{i}. {ref}")
    save_doc(doc, "IJMR_frailty_blinded_manuscript_scaffold_2026-05-18.docx")

    ref_df = pd.DataFrame(
        [
            [1, "Page MJ", "PRISMA 2020", "10.1136/bmj.n71", "https://www.prisma-statement.org/prisma-2020"],
            [2, "Hutton B", "PRISMA-NMA extension", "10.7326/M14-2385", "https://www.prisma-statement.org/nma"],
            [3, "Indian Journal of Medical Research", "Instructions for authors", "", "https://ijmr.org.in/for-authors/"],
            [4, "Cochrane Methods", "Risk of Bias 2 tool", "", "https://methods.cochrane.org/risk-bias-2"],
        ],
        columns=["citation_number", "first_author_or_source", "topic", "doi", "url"],
    )
    ref_df.to_csv(ASSET_DIR / "method_reference_seed_metadata.csv", index=False)


def build_cover_letter_hold() -> None:
    doc = Document()
    set_normal_style(doc)
    add_heading(doc, "Cover Letter Draft - Hold Until Evidence Gates Complete")
    doc.add_paragraph("To\nThe Editor\nIndian Journal of Medical Research")
    doc.add_paragraph(
        "Please find enclosed a manuscript entitled \""
        + TITLE
        + "\" by "
        + AUTHOR_1
        + " and "
        + AUTHOR_2
        + "."
    )
    doc.add_paragraph(
        "This draft cover letter is not for immediate submission. The systematic review requires protocol registration, final full-text screening, risk-of-bias assessment and synthesis before portal upload."
    )
    doc.add_paragraph(
        "The proposed manuscript is intended to address a nationally relevant question: which community-deliverable frailty interventions are both effective and feasible for Indian primary care."
    )
    doc.add_paragraph("Sincerely,\n" + AUTHOR_1 + "\nCorresponding author details pending")
    save_doc(doc, "IJMR_frailty_cover_letter_hold_2026-05-18.docx")


def build_asset_index() -> None:
    lines = [
        "# IJMR Frailty Interventions Audit-Ready Asset Index",
        "",
        "Generated: 2026-05-18",
        "",
        "## Submission Components",
        "",
        "- `IJMR_frailty_first_page_2026-05-18.docx`",
        "- `IJMR_frailty_declarations_2026-05-18.docx`",
        "- `IJMR_frailty_blinded_manuscript_scaffold_2026-05-18.docx`",
        "- `IJMR_frailty_figures_2026-05-18.docx`",
        "- `IJMR_frailty_supplementary_2026-05-18.docx`",
        "- `IJMR_frailty_cover_letter_hold_2026-05-18.docx`",
        "- `submission_readiness_audit_2026-05-18.md`",
        "- `submission_readiness_audit.csv`",
        "- `internal_double_blind_peer_review_2026-05-18.md`",
        "- `author_response_revision_plan_2026-05-18.md`",
        "- `method_reference_seed_metadata.csv`",
        "",
        "## Figures",
        "",
        "- `figures/figure1_screening_progress_flow.png`",
        "- `figures/figure2_preliminary_node_distribution.png`",
        "",
        "## Readiness Decision",
        "",
        "Not ready for journal submission. The assets are structurally prepared for audit and author completion, but final submission requires protocol registration, final full-text inclusion/exclusion, numeric extraction, RoB 2 assessment, final reference list, and synthesis.",
    ]
    (ASSET_DIR / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig1 = make_prisma_progress_figure()
    fig2 = make_node_figure()
    build_first_page()
    build_declarations()
    build_figures_docx(fig1, fig2)
    build_supplementary()
    build_manuscript_scaffold()
    build_cover_letter_hold()
    build_asset_index()
    print(f"submission_assets={ASSET_DIR}")
    for path in sorted(ASSET_DIR.rglob("*")):
        if path.is_file():
            print(path.relative_to(ASSET_DIR))


if __name__ == "__main__":
    main()
