"""Build author data-request and manual rescue package for non-extractable DNAm studies."""
from __future__ import annotations
import argparse
import sys
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt
sys.path.insert(0, str(Path(__file__).parent))
from _common import ensure_dirs, load_config, log  # noqa: E402


REQUEST_FIELDS = [
    "study_id", "title", "doi", "pmid", "corresponding_author_email", "clock_version",
    "tissue_or_cell_type", "array_platform", "arm_name", "n_baseline", "n_followup",
    "baseline_mean", "baseline_sd", "followup_mean", "followup_sd",
    "change_mean", "change_sd", "timepoint_weeks", "notes",
]


def style(doc):
    s = doc.styles["Normal"]; s.font.name = "Times New Roman"; s.font.size = Pt(12)


def run(cfg):
    freeze = cfg["project"]["freeze_date"]
    proc = Path(cfg["paths"]["data_processed"]); tabs = Path(cfg["paths"]["results_tabs"]); docs_dir = Path(cfg["paths"]["docs"]); sub = Path(cfg["paths"]["submission"])
    ensure_dirs(tabs, docs_dir, sub)
    audit = pd.read_csv(proc / f"fulltext_eligibility_audit_{freeze}.csv").fillna("")
    extracted = pd.read_csv(proc / f"extracted_clock_studies_{freeze}.csv").fillna("")
    effects = pd.read_csv(tabs / f"effect_size_candidates_{freeze}.csv").fillna("")

    non_extract = extracted[~extracted["study_id"].isin(effects[effects["effect_status"].eq("calculated_candidate")]["study_id"])].copy()
    priority_ids = set(audit[audit["first_reviewer_fulltext_decision"].isin(["include_accessible_first_reviewer", "await_fulltext"])]["study_id"])
    non_extract = non_extract[non_extract["study_id"].isin(priority_ids)].copy()

    rescue_rows = []
    for sid, group in non_extract.groupby("study_id"):
        audit_row = audit[audit["study_id"].eq(sid)].head(1)
        decision = audit_row["first_reviewer_fulltext_decision"].iloc[0] if not audit_row.empty else ""
        reason = audit_row["first_reviewer_reason"].iloc[0] if not audit_row.empty else ""
        clocks = ";".join(sorted(set(group["clock"].astype(str))))
        rescue_rows.append({
            "study_id": sid,
            "title": audit_row["title"].iloc[0] if not audit_row.empty else "",
            "doi": group["doi"].iloc[0],
            "pmid": group["pmid"].iloc[0],
            "first_reviewer_fulltext_decision": decision,
            "first_reviewer_reason": reason,
            "clocks_seeded": clocks,
            "extraction_statuses": ";".join(sorted(set(group["extraction_status"].astype(str)))),
            "priority": "high" if decision == "include_accessible_first_reviewer" else "medium",
            "recommended_action": "manual supplement/figure extraction and author data request" if decision == "include_accessible_first_reviewer" else "retrieve full text then assess",
        })
    rescue = pd.DataFrame(rescue_rows).sort_values(["priority", "study_id"], ascending=[True, True])
    rescue.to_csv(tabs / f"data_rescue_priority_{freeze}.csv", index=False)

    template = []
    for _, row in rescue.iterrows():
        for arm in ["intervention", "control"]:
            item = {field: "" for field in REQUEST_FIELDS}
            item.update({
                "study_id": row["study_id"], "title": row["title"], "doi": row["doi"], "pmid": row["pmid"], "arm_name": arm
            })
            template.append(item)
    pd.DataFrame(template).to_csv(tabs / f"author_data_request_template_{freeze}.csv", index=False)

    doc = Document(); style(doc)
    doc.add_heading("Author Data Request Template", 0)
    doc.add_paragraph("Dear Dr [Corresponding author],")
    doc.add_paragraph(
        "We are conducting a systematic review of human intervention studies reporting DNA-methylation ageing clocks. "
        "Your study appears potentially eligible, but the published report does not provide all arm-level values required for transparent effect-size calculation."
    )
    doc.add_paragraph(
        "Could you please share, for each randomized arm and DNAm clock, the sample size, baseline mean and SD, follow-up mean and SD, "
        "change-score mean and SD, timepoint, tissue or cell type, clock version, methylation array and preprocessing details?"
    )
    doc.add_paragraph(
        "We will use the data only for aggregate evidence synthesis and will cite the original publication. If data cannot be shared, a brief confirmation would still help us document availability transparently."
    )
    doc.add_paragraph("Sincerely,\nDr Siddalingaiah H S\nDr Chandrakala D")
    doc.save(sub / f"DNAm_clock_author_data_request_template_{freeze}.docx")

    report = "# Data Rescue Package\n\n"
    report += f"Generated: {freeze}\n\n"
    report += f"Priority records requiring manual rescue or author data request: {len(rescue)}\n\n"
    report += rescue.to_markdown(index=False)
    report += "\n\nFiles created:\n\n"
    report += f"- `results/tables/data_rescue_priority_{freeze}.csv`\n"
    report += f"- `results/tables/author_data_request_template_{freeze}.csv`\n"
    report += f"- `submission_assets/DNAm_clock_author_data_request_template_{freeze}.docx`\n"
    (docs_dir / f"data_rescue_package_{freeze}.md").write_text(report, encoding="utf-8")
    log("data_rescue_package_done", priority_records=len(rescue))


if __name__ == "__main__":
    ap = argparse.ArgumentParser(); ap.add_argument("--config", required=True)
    run(load_config(ap.parse_args().config))
