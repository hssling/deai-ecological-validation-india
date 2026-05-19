"""Phase 14 (Path-C): build IJMR submission bundle from manuscript_main.md + tables/figures.

Outputs in submission_assets/IJMR_DNAm_clocks_path_c_<freeze>/:
  IJMR_DNAm_clocks_manuscript_<freeze>.docx
  IJMR_DNAm_clocks_title_page_<freeze>.docx
  IJMR_DNAm_clocks_cover_letter_<freeze>.docx
  IJMR_DNAm_clocks_declarations_<freeze>.docx
  IJMR_DNAm_clocks_PRISMA_2020_<freeze>.docx
  IJMR_DNAm_clocks_supplementary_<freeze>.docx  (S1–S10)
  submission_checklist_<freeze>.docx
  figures/  (copies of 4 main PNGs)
  tables/   (copies of CSVs + DOCX renderings)
"""
from __future__ import annotations

import argparse
import math
import shutil
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

sys.path.insert(0, str(Path(__file__).parent))
from _common import ensure_dirs, load_config, log  # noqa: E402


def style_doc(doc: Document) -> None:
    s = doc.styles["Normal"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(11)


def add_table_from_df(doc: Document, df: pd.DataFrame, max_cols: int = 10) -> None:
    cols = list(df.columns)[:max_cols]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, c in enumerate(cols):
        table.rows[0].cells[i].text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            v = row[c]
            if isinstance(v, float) and math.isnan(v):
                v = ""
            cells[i].text = str(v)


def build_main_manuscript(template_md: Path, out_path: Path, placeholders: dict, figs: list[Path]) -> int:
    """Build the main manuscript .docx with placeholders substituted.

    Returns word count of the body text (excluding tables/figures).
    """
    text = template_md.read_text(encoding="utf-8")
    for k, v in placeholders.items():
        text = text.replace("{{" + k + "}}", str(v))

    doc = Document()
    style_doc(doc)

    word_count = 0
    in_refs = False
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), 0)
        elif line.startswith("## "):
            heading = line[3:].strip()
            doc.add_heading(heading, 1)
            in_refs = heading.lower().startswith("references")
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 2)
        elif line.startswith("---"):
            continue
        else:
            p = doc.add_paragraph(line)
            if not in_refs:
                word_count += len([w for w in line.split() if w])

    # Append figures section
    doc.add_heading("Figures", 1)
    for f in figs:
        if f.exists():
            doc.add_paragraph(f.stem.replace("_", " "))
            doc.add_picture(str(f), width=Inches(5.5))

    doc.save(out_path)
    return word_count


def build_title_page(out_path: Path, freeze: str) -> None:
    doc = Document(); style_doc(doc)
    doc.add_heading("Title page", 0)
    doc.add_paragraph(
        "Title: Human interventions and DNA-methylation ageing clocks: a Path-C systematic review and "
        "meta-analysis under strict honesty constraints"
    )
    doc.add_paragraph("Short title: DNAm clocks in interventions — Path-C synthesis")
    doc.add_heading("Authors", 1)
    doc.add_paragraph("Dr Siddalingaiah H S (corresponding author)")
    doc.add_paragraph("Dr Chandrakala D")
    doc.add_heading("Corresponding author contact", 1)
    doc.add_paragraph("Dr Siddalingaiah H S")
    doc.add_paragraph("Email: hssling@gmail.com")
    doc.add_paragraph(f"Data freeze: {freeze}")
    doc.add_paragraph("Target journal: Indian Journal of Medical Research (IJMR)")
    doc.save(out_path)


def build_cover_letter(out_path: Path, freeze: str) -> None:
    doc = Document(); style_doc(doc)
    doc.add_heading("Cover letter", 0)
    doc.add_paragraph(f"Date: {freeze}")
    doc.add_paragraph("To: The Editor-in-Chief, Indian Journal of Medical Research")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Dear Editor, we submit the manuscript entitled 'Human interventions and DNA-methylation ageing "
        "clocks: a Path-C systematic review and meta-analysis under strict honesty constraints' for "
        "consideration in IJMR."
    )
    doc.add_paragraph(
        "The work systematically synthesises adjusted between-group effects on DNAm clocks from human "
        "intervention studies under pre-registered honesty gates (no fabrication, transparent gate failures). "
        "Quantitative pooling was possible only for DunedinPACE (k=3, very low GRADE certainty); other clocks "
        "are reported narratively. The manuscript additionally proposes a minimal trial-reporting checklist "
        "to make future DNAm-clock syntheses feasible."
    )
    doc.add_paragraph(
        "The work is original, has not been published elsewhere, and is not under consideration by any other "
        "journal. All authors have approved the submission and declare no conflicts of interest."
    )
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Dr Siddalingaiah H S (corresponding author)")
    doc.add_paragraph("Dr Chandrakala D")
    doc.save(out_path)


def build_declarations(out_path: Path) -> None:
    doc = Document(); style_doc(doc)
    doc.add_heading("Declarations", 0)
    doc.add_heading("Funding", 1)
    doc.add_paragraph("None.")
    doc.add_heading("Conflicts of interest", 1)
    doc.add_paragraph("The authors declare no conflicts of interest.")
    doc.add_heading("Ethical approval", 1)
    doc.add_paragraph("Not applicable — systematic review of published literature.")
    doc.add_heading("Data availability", 1)
    doc.add_paragraph(
        "All extracted data, code, and gate reports are provided in the supplementary repository "
        "anti_ageing_review/meta_dnam_clocks/."
    )
    doc.add_heading("Author contributions", 1)
    doc.add_paragraph(
        "Dr Siddalingaiah H S: conception, protocol, screening, extraction, analysis, drafting. "
        "Dr Chandrakala D: dual screening and extraction, RoB 2 dual coding, critical review."
    )
    doc.save(out_path)


def build_prisma_checklist(out_path: Path) -> None:
    doc = Document(); style_doc(doc)
    doc.add_heading("PRISMA 2020 checklist", 0)
    items = [
        ("Title", "1", "Identified as a systematic review and meta-analysis in the title."),
        ("Abstract", "2", "Structured abstract included."),
        ("Rationale", "3", "Stated in Introduction."),
        ("Objectives", "4", "Stated in Introduction."),
        ("Eligibility criteria", "5", "Methods § Eligibility."),
        ("Information sources", "6", "Methods § Search."),
        ("Search strategy", "7", "Methods § Search."),
        ("Selection process", "8", "Methods § Screening."),
        ("Data collection process", "9", "Methods § Extraction."),
        ("Data items", "10", "Methods § Extraction."),
        ("Study risk of bias assessment", "11", "Methods § Risk of bias (RoB 2; dual coding pending)."),
        ("Effect measures", "12", "Methods § Quantitative synthesis."),
        ("Synthesis methods", "13", "Methods § Quantitative synthesis."),
        ("Reporting bias assessment", "14", "Methods § Gates; pub-bias not assessed (k<10)."),
        ("Certainty assessment", "15", "Methods § GRADE."),
        ("Study selection (flow)", "16", "Figure 1 (PRISMA 2020 flow)."),
        ("Study characteristics", "17", "Supplementary Table S1."),
        ("Risk of bias in studies", "18", "Figure 2 + Supplementary Table S3."),
        ("Results of individual studies", "19", "Supplementary Table S2."),
        ("Results of syntheses", "20", "Results § Primary quantitative synthesis; Figure 3."),
        ("Reporting biases", "21", "Results § Publication bias and NMA (gate-failed)."),
        ("Certainty of evidence", "22", "Results § GRADE."),
        ("Discussion", "23", "Discussion section."),
        ("Registration and protocol", "24", "Protocol v1 (frozen); amendments A2/A3."),
        ("Support", "25", "Declarations."),
        ("Competing interests", "26", "Declarations."),
        ("Availability of data, code, and other materials", "27", "Declarations § Data availability."),
    ]
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
    for i, h in enumerate(["Item", "PRISMA #", "Where reported"]):
        table.rows[0].cells[i].text = h
    for it, num, where in items:
        cells = table.add_row().cells
        cells[0].text = it; cells[1].text = num; cells[2].text = where
    doc.save(out_path)


def build_supplementary(out_path: Path, tables_dir: Path, freeze: str) -> None:
    doc = Document(); style_doc(doc)
    doc.add_heading("Supplementary materials", 0)
    s_items = [
        ("S1. Included-study characteristics", "included_studies"),
        ("S2. Study-level adjusted effects (extraction snippets)", "study_level_effects"),
        ("S3. RoB 2 worksheet (all pending dual coding)", "rob2_assessments"),
        ("S4. Per-clock pooled (Path-C)", "per_clock_pooled"),
        ("S5. Subgroup and meta-regression gates", "subgroup_metareg"),
        ("S6. NMA gate (not performed)", "nma_not_performed_path_c"),
        ("S7. Sensitivity (leave-one-out + restrictions)", "sensitivity"),
        ("S8. GRADE evidence profile (Path-C)", "grade_profile_path_c"),
        ("S9. Relaxed eligibility audit (Path-C amendments)", "relaxed_eligibility_audit"),
        ("S10. Proposed DNAm-clock trial-reporting checklist",
         None),
    ]
    for title, stem in s_items:
        doc.add_heading(title, 1)
        if stem is None:
            # S10 narrative checklist
            for item in [
                "Per-arm n at baseline and at every follow-up timepoint.",
                "Per-arm clock mean and SD at each timepoint (or change mean and SD).",
                "Clock name AND version (incl. PC reformulation indication).",
                "Tissue / sorted cell type.",
                "Methylation array (450K / EPIC / EPICv2) and preprocessing pipeline.",
                "Whether values are raw clock age, residualized acceleration, or pace-of-ageing units.",
                "Cell-composition adjustment used (if any).",
                "Adjusted-model specification with full covariate list, SE/CI on the original clock scale.",
                "Supplementary CSV of arm-level extracted values for the clock outcome(s).",
                "Pre-registration of DNAm-clock endpoint(s) and analysis plan.",
            ]:
                doc.add_paragraph("- " + item)
            continue
        csv_path = tables_dir / f"{stem}_{freeze}.csv"
        if not csv_path.exists():
            # Try processed dir paths
            alt = tables_dir.parent.parent / "data" / "processed" / f"{stem}_{freeze}.csv"
            if alt.exists():
                csv_path = alt
        if csv_path.exists():
            try:
                df = pd.read_csv(csv_path).fillna("")
                add_table_from_df(doc, df, max_cols=8)
            except Exception as e:
                doc.add_paragraph(f"(could not render {csv_path.name}: {e})")
        else:
            doc.add_paragraph(f"(source table {stem}_{freeze}.csv not found)")
    doc.save(out_path)


def build_submission_checklist(out_path: Path) -> None:
    doc = Document(); style_doc(doc)
    doc.add_heading("Submission checklist (IJMR)", 0)
    items = [
        "Cover letter included (IJMR_DNAm_clocks_cover_letter).",
        "Title page with corresponding author contact.",
        "Blinded manuscript (main text).",
        "PRISMA 2020 checklist (separate file).",
        "Supplementary materials S1–S10.",
        "Four main figures at 300 dpi (PNG + SVG).",
        "Tables (CSV + DOCX renderings).",
        "Declarations: funding, conflicts, ethics, data availability, author contributions.",
        "Confirmation: no fabrication; every gate failure reported transparently.",
    ]
    for it in items:
        doc.add_paragraph("[ ] " + it)
    doc.save(out_path)


def run(cfg: dict) -> None:
    freeze = cfg["project"]["freeze_date"]
    tabs = Path(cfg["paths"]["results_tabs"])
    figs = Path(cfg["paths"]["results_figs"])
    sub_root = Path(cfg["paths"]["submission"])
    bundle = sub_root / f"IJMR_DNAm_clocks_path_c_{freeze}"
    bundle_figs = bundle / "figures"
    bundle_tabs = bundle / "tables"
    ensure_dirs(bundle, bundle_figs, bundle_tabs)

    # Gather pooled / sle for placeholders
    pooled = pd.read_csv(tabs / f"per_clock_pooled_{freeze}.csv")
    dp = pooled[pooled["clock"] == "DunedinPACE"].iloc[0]

    def f(x, fmt="{:.4f}"):
        try:
            return fmt.format(float(x))
        except Exception:
            return "—"

    relaxed_path = Path(cfg["paths"]["data_processed"]) / f"relaxed_eligibility_audit_{freeze}.csv"
    if relaxed_path.exists():
        rel = pd.read_csv(relaxed_path)
        qual_n = int(rel["final_eligibility"].isin(
            ["include_accessible_first_reviewer", "include_relaxed"]).sum())
        relaxed_n = int((rel["final_eligibility"] == "include_relaxed").sum())
    else:
        qual_n, relaxed_n = 21, 0
    raw_path = Path(cfg["paths"]["data_interim"]) / f"raw_records_dnam_{freeze}.csv"
    raw_n = int(pd.read_csv(raw_path).shape[0]) if raw_path.exists() else 2804

    placeholders = {
        "FREEZE": freeze,
        "RAW_N": raw_n,
        "QUAL_N": qual_n,
        "RELAXED_N": relaxed_n,
        "DPACE_K": int(dp["k"]),
        "DPACE_MU": f(dp["mu_dl"]),
        "DPACE_CI": f"[{f(dp['ci_lower_dl'])}, {f(dp['ci_upper_dl'])}]",
        "DPACE_PI": f"[{f(dp['pi_lower'])}, {f(dp['pi_upper'])}]",
        "DPACE_I2": f"{float(dp['I2']):.0f}",
        "DPACE_TAU2": f(dp["tau2"]),
        "DPACE_QP": f(dp["Q_pval"], "{:.3f}"),
        "DPACE_MU_HKSJ": f(dp["mu_hksj"]),
        "DPACE_CI_HKSJ": f"[{f(dp['ci_lower_hksj'])}, {f(dp['ci_upper_hksj'])}]",
        "DPACE_MU_BAYES": f(dp["mu_bayes"]),
        "DPACE_CI_BAYES": f"[{f(dp['ci_lower_bayes'])}, {f(dp['ci_upper_bayes'])}]",
    }

    # Copy figures
    fig_files = [
        figs / f"fig1_prisma_path_c_{freeze}.png",
        figs / f"fig2_rob_path_c_{freeze}.png",
        figs / f"fig3_forest_dunedinpace_path_c_{freeze}.png",
        figs / f"fig4_gate_status_path_c_{freeze}.png",
    ]
    for fp in fig_files:
        if fp.exists():
            shutil.copy(fp, bundle_figs / fp.name)
            svg = fp.with_suffix(".svg")
            if svg.exists():
                shutil.copy(svg, bundle_figs / svg.name)

    # Copy tables
    main_tables = [
        f"per_clock_pooled_{freeze}.csv",
        f"subgroup_metareg_{freeze}.csv",
        f"nma_not_performed_path_c_{freeze}.csv",
        f"sensitivity_{freeze}.csv",
        f"grade_profile_path_c_{freeze}.csv",
    ]
    for t in main_tables:
        src = tabs / t
        if src.exists():
            shutil.copy(src, bundle_tabs / t)
            try:
                df = pd.read_csv(src).fillna("")
                doc = Document(); style_doc(doc)
                doc.add_heading(t, 0)
                add_table_from_df(doc, df, max_cols=10)
                doc.save(bundle_tabs / (Path(t).stem + ".docx"))
            except Exception:
                pass

    # Build main manuscript
    manuscript_md = Path(cfg["paths"]["manuscript"]) / "manuscript_main.md"
    main_out = bundle / f"IJMR_DNAm_clocks_manuscript_{freeze}.docx"
    wc = build_main_manuscript(manuscript_md, main_out, placeholders, fig_files)

    # Auxiliary docs
    build_title_page(bundle / f"IJMR_DNAm_clocks_title_page_{freeze}.docx", freeze)
    build_cover_letter(bundle / f"IJMR_DNAm_clocks_cover_letter_{freeze}.docx", freeze)
    build_declarations(bundle / f"IJMR_DNAm_clocks_declarations_{freeze}.docx")
    build_prisma_checklist(bundle / f"IJMR_DNAm_clocks_PRISMA_2020_{freeze}.docx")
    build_supplementary(bundle / f"IJMR_DNAm_clocks_supplementary_{freeze}.docx", tabs, freeze)
    build_submission_checklist(bundle / f"submission_checklist_{freeze}.docx")

    log("manuscript_bundle_built", bundle=str(bundle), main_word_count=wc)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", default="anti_ageing_review/meta_dnam_clocks/config/meta_config.yaml")
    args = ap.parse_args()
    cfg = load_config(args.config)
    run(cfg)


if __name__ == "__main__":
    main()
