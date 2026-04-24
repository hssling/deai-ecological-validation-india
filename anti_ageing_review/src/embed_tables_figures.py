"""
Rebuild the corrected blinded manuscript DOCX with:
  - Embedded real tables (from CSV data)
  - Embedded figures (PNG images)
  - Superscript square-bracket citations  e.g. [1,2]
  - Updated Data Availability Statement with GitHub URL
Run from D:/Anti ageing research directory.
"""

import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = "anti_ageing_review/submission_assets/MJDRDYPU_AntiAgeing_final_qa_2026-04-24"
TABLES = f"{BASE}/tables"
FIGURES = f"{BASE}/figures"
OUT    = f"{BASE}/MJDRDYPU_AntiAgeing_blinded_manuscript_embedded_2026-04-24.docx"

GITHUB_URL = "https://github.com/hssling/deai-ecological-validation-india"

# Regex: matches [1], [1,2], [1-14], [15-40], [3,4,5,6,7,8,9,10] etc.
_CIT_RE = re.compile(r'(\[[\d,\-]+\])')


# ── helpers ───────────────────────────────────────────────────────────

def set_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    for s in doc.sections:
        s.top_margin    = Inches(top)
        s.bottom_margin = Inches(bottom)
        s.left_margin   = Inches(left)
        s.right_margin  = Inches(right)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def _add_runs(p, text, bold=False, italic=False, size=11):
    """Add text to paragraph p, making [n] citation tokens superscript."""
    tokens = _CIT_RE.split(text)
    for token in tokens:
        r = p.add_run(token)
        r.bold        = bold
        r.italic      = italic
        r.font.size   = Pt(size)
        if _CIT_RE.fullmatch(token):
            r.font.superscript = True


def para(doc, text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    _add_runs(p, text, bold=bold, italic=italic, size=size)
    if align:
        p.alignment = align
    return p


def bold_label(doc, label, text, size=11):
    """Bold label + body text, with superscript citations in the body."""
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(size)
    _add_runs(p, text, size=size)
    return p


def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


def shade_row(row, hex_color="D9E1F2"):
    """Light blue header shading."""
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)


def bold_row(row, size=9):
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(size)


def fill_row(row, values, size=9):
    for i, v in enumerate(values):
        if i < len(row.cells):
            row.cells[i].text = str(v) if not pd.isna(v) else ""
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(size)


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    return p


def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    return p


def embed_figure(doc, path, width_inches=6.0, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        add_figure_caption(doc, caption)


# ── label cleaners ────────────────────────────────────────────────────

INTERV_LABELS = {
    "exercise":            "Exercise",
    "microbiome":          "Microbiome",
    "rapamycin_mtor":      "Rapamycin/mTOR",
    "senolytics":          "Senolytics",
    "caloric_restriction": "Caloric restriction",
    "lifestyle_bundle":    "Lifestyle bundle",
    "nad_sirtuin":         "NAD+/sirtuin",
    "fasting":             "Fasting",
    "supplements":         "Supplements",
    "metformin":           "Metformin",
    "sleep_circadian":     "Sleep/circadian",
    "controversial":       "Controversial",
    "reprogramming":       "Reprogramming",
    "stem_cell":           "Stem cell",
}

TRANS_LABELS = {
    "A_healthspan_support_signal":          "A – Healthspan support signal",
    "B_promising_not_recommendation_ready": "B – Promising; not recommendation-ready",
    "C_biomarker_or_indirect":              "C – Biomarker / indirect signal",
    "D_speculative_or_low_directness":      "D – Speculative / low directness",
}

DOMAIN_LABELS = {
    "healthspan_functional_ageing":  "Healthspan / function",
    "hard_ageing_relevance":         "Hard ageing relevance",
    "biological_ageing_biomarker":   "Biological-age biomarker",
    "surrogate_or_indirect":         "Surrogate / indirect",
}

ROB_LABELS = {
    "low_concern":                                "Low",
    "some_concern":                               "Some concern",
    "unclear":                                    "Unclear",
    "low_to_some_concern_pending_full_text_review": "Low–Some (pending FT)",
    "serious_or_unclear":                         "Serious/unclear",
    "low_concern_rct_design":                     "Low (RCT design)",
    "some_concern_observational_confounding":     "Some (observational)",
    "not_applicable":                             "NA",
    "not_applicable_or_unclear":                  "NA / unclear",
    "unclear_without_protocol":                   "Unclear (no protocol)",
    "unclear_or_not_applicable":                  "Unclear / NA",
    "low_or_some_concern":                        "Low–Some",
}

# Curated effect summaries for refs 1–14 (extracted from candidate_effect_size_text)
EFFECT_SUMMARY = {
    0:  "Baseline characteristics paper (n=1212 randomised, mean age 68.7 yr). "
        "No primary outcome effect data reported yet.",
    1:  "n=16 (10 exercise, 6 control). Frailty Index: ΔEFIP −0.07 (95% CI −0.14 to −0.00). "
        "Mental health QoL: +21.24 (95% CI 7.32 to 35.16). Adherence 97.2%.",
    2:  "Chair rise (primary): β=0.579 (95% CI −1.08 to 2.24; p=0.494; NS). "
        "Cognition: β=−0.482 (95% CI −0.813 to −0.141; p=0.014, favours prebiotic).",
    3:  "n=60 randomised (55 completed). Significant group-by-time improvements in "
        "handgrip strength and most frailty/sarcopenia outcomes (p-values not re-extracted).",
    4:  "Full text not retrieved. Effect data not extracted in present synthesis.",
    5:  "5 cohorts (SHARE n=56,555; CHARLS n=12,271 + 3 others). PA modified "
        "frailty–mortality association. Hazard ratios not re-extracted in present synthesis.",
    6:  "3 FMD cycles: median biological-age reduction of 2.5 yr (independent of "
        "weight loss). Hepatic biomarkers and disease-risk markers improved.",
    7:  "p16INK4A reduction (p=0.008); collagen VII increase (significant). "
        "Skin-ageing senescence markers reduced by topical rapamycin.",
    8:  "Lean tissue mass, pain, and general health improved. "
        "Visceral adiposity: ηp²=0.001, p=0.942 (NS). Blood biomarkers within normal ranges.",
    9:  "Abstract only. No primary numeric outcome data re-extracted in present synthesis.",
    10: "Full text not retrieved. Effect data not extracted in present synthesis.",
    11: "3-arm RCT (CR in-person; CR remote; TRE). Primary outcomes: feasibility "
        "and adherence. Clinical effect data pending final extraction.",
    12: "n=47 frail adults (mean age 80.2 yr, SD 3.1). Statistically significant "
        "improvements in functional trajectories and epigenetic ageing markers vs control.",
    13: "Trial design / eligibility criteria paper (TROFFi phase II fisetin RCT). "
        "No outcome effect data available.",
}

# Short titles for tables
SHORT_TITLES = {
    0:  "SINGER multidomain RCT — baseline characteristics",
    1:  "Multimodal exercise in frail MS patients (6-week RCT)",
    2:  "PROMOTe microbiome RCT — muscle function and cognition",
    3:  "Ring Fit exergame RCT for frailty/sarcopenia",
    4:  "Resistance exercise in cognitive frailty RCT",
    5:  "Physical activity & deficit frailty: 5-cohort mortality study",
    6:  "Fasting-mimicking diet and biological-age markers (RCT)",
    7:  "Topical rapamycin and skin senescence markers (RCT)",
    8:  "PEARL rapamycin trial — 1-year healthspan metrics",
    9:  "FMD: multi-system regeneration and healthspan (2015 RCT)",
    10: "Creatine + resistance training in frail older subjects",
    11: "HALLO pilot: caloric restriction and TRE in older adults",
    12: "Multidomain lifestyle intervention & epigenetic ageing (RCT)",
    13: "TROFFi fisetin phase II RCT design (breast cancer survivors)",
}

REF_NUMBERS = {i: i + 1 for i in range(14)}  # row 0 → Ref 1 … row 13 → Ref 14


# ── Table 1: intervention credibility ranking ─────────────────────────

def build_table1(doc):
    df = pd.read_csv(f"{TABLES}/intervention_credibility_ranking.csv")

    add_table_caption(
        doc,
        "Table 1. Intervention credibility ranking derived from extracted evidence "
        "(n=14 intervention classes; 1029 deduplicated records). Credibility score "
        "incorporates human record count, human trial count, direct ageing/healthspan "
        "outcomes, biomarker evidence, surrogate burden, and hype-language penalty. "
        "Rankings are provisional pending final full-text adjudication."
    )

    cols = [
        "Rank", "Intervention class", "Records\n(n)",
        "Human\nrecords", "Human\ntrials",
        "Credibility\nscore", "Translational category",
    ]
    widths = [0.45, 1.15, 0.65, 0.65, 0.65, 0.75, 2.15]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(cols, widths)):
        hdr.cells[i].text = h
        hdr.cells[i].width = Inches(w)
        for p in hdr.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    shade_row(hdr)

    for _, row_data in df.iterrows():
        row = table.add_row()
        values = [
            int(row_data["rank"]),
            INTERV_LABELS.get(row_data["intervention_name"], row_data["intervention_name"]),
            int(row_data["n_extracted_records"]),
            int(row_data["human_records"]),
            int(row_data["human_trial_records"]),
            f"{row_data['credibility_score']:.2f}",
            TRANS_LABELS.get(row_data["credibility_tier"], row_data["credibility_tier"]),
        ]
        for i, (v, w) in enumerate(zip(values, widths)):
            row.cells[i].text = str(v)
            row.cells[i].width = Inches(w)
            for p in row.cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 and i != 6 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()


# ── Table 2: priority human evidence + effect estimates ──────────────

def build_table2(doc):
    df = pd.read_csv(f"{TABLES}/effect_size_extraction_priority_human.csv")
    df14 = df.head(14).copy()

    add_table_caption(
        doc,
        "Table 2. Representative priority human evidence with effect estimates extracted "
        "from available open text (n=14 records; refs 1–14). Effect information is "
        "retained as extracted text; estimates are not harmonised or pooled. "
        "Records labelled 'full text not retrieved' require manual extraction."
    )

    cols = ["Ref", "Study (abbreviated)", "Intervention", "Outcome domain", "Key extracted estimate"]
    widths = [0.35, 1.70, 0.95, 1.05, 2.45]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(cols, widths)):
        hdr.cells[i].text = h
        hdr.cells[i].width = Inches(w)
        for p in hdr.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    shade_row(hdr)

    for i, row_data in df14.iterrows():
        row = table.add_row()
        values = [
            str(i + 1),
            SHORT_TITLES.get(i, str(row_data["title"])[:60]),
            INTERV_LABELS.get(row_data["intervention_name"], row_data["intervention_name"]),
            DOMAIN_LABELS.get(row_data["ageing_domain_category"], row_data["ageing_domain_category"]),
            EFFECT_SUMMARY.get(i, "See source text."),
        ]
        for j, (v, w) in enumerate(zip(values, widths)):
            row.cells[j].text = str(v)
            row.cells[j].width = Inches(w)
            for p in row.cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.add_paragraph()


# ── Table 3: preliminary risk of bias ────────────────────────────────

def build_table3(doc):
    df = pd.read_csv(f"{TABLES}/risk_of_bias_formal_preliminary_human.csv")
    df14 = df.head(14).copy()

    add_table_caption(
        doc,
        "Table 3. Preliminary risk-of-bias domain prompts for priority human records "
        "(n=14; refs 1–14). These are structured prompts based on available open text or "
        "abstract content; they are NOT final RoB 2, ROBINS-I, SYRCLE, or GRADE "
        "assessments. Final judgements require complete full-text review."
    )

    cols = [
        "Ref", "Study (abbreviated)",
        "Randomisation", "Blinding", "Missing data",
        "Confounding", "Overall (preliminary)",
    ]
    widths = [0.35, 1.60, 0.80, 0.70, 0.75, 0.90, 1.40]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(cols, widths)):
        hdr.cells[i].text = h
        hdr.cells[i].width = Inches(w)
        for p in hdr.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    shade_row(hdr)

    for i, row_data in df14.iterrows():
        row = table.add_row()
        values = [
            str(i + 1),
            SHORT_TITLES.get(i, str(row_data["title"])[:55]),
            ROB_LABELS.get(row_data["rob_randomization"],   row_data["rob_randomization"]),
            ROB_LABELS.get(row_data["rob_blinding"],         row_data["rob_blinding"]),
            ROB_LABELS.get(row_data["rob_missing_data"],     row_data["rob_missing_data"]),
            ROB_LABELS.get(row_data["rob_confounding"],      row_data["rob_confounding"]),
            ROB_LABELS.get(row_data["rob_overall"],          row_data["rob_overall"]),
        ]
        for j, (v, w) in enumerate(zip(values, widths)):
            row.cells[j].text = str(v)
            row.cells[j].width = Inches(w)
            for p in row.cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.add_paragraph()


# ── main: build full embedded manuscript ──────────────────────────────

def build_manuscript():
    doc = Document()
    set_margins(doc)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # ── Title block ────────────────────────────────────────────────
    para(doc,
         "Can Ageing Be Slowed or Reversed? A Reproducible Evidence Map, "
         "Credibility Ranking, and Mechanistic Synthesis of Anti-Ageing and "
         "Age-Reversal Interventions",
         bold=True, size=14)
    para(doc, "Article type: Systematic review / evidence map", size=11)
    para(doc, "Running title: Anti-ageing evidence map", size=11)
    doc.add_paragraph()

    # ── Abstract ───────────────────────────────────────────────────
    heading(doc, "Abstract", 1)
    bold_label(doc, "Background: ",
               "Anti-ageing claims range from established healthspan interventions to "
               "speculative rejuvenation approaches.")
    bold_label(doc, "Objective: ",
               "To create a conservative, reproducible evidence map separating healthspan "
               "benefit, lifespan extension, biological-age slowing, biomarker reversal, "
               "and true clinical rejuvenation.")
    bold_label(doc, "Methods: ",
               "Searches across PubMed, Europe PMC, and Crossref retrieved 1155 raw records, "
               "deduplicated to 1029. Title/abstract screening, metadata-assisted extraction, "
               "credibility scoring, mechanism mapping, priority human verification, preliminary "
               "risk-of-bias assessment, duplicate checks, and effect-estimate extraction were "
               "performed.")
    bold_label(doc, "Results: ",
               "Screening classified 29 records as include, 455 as uncertain, and 545 as "
               "exclude. Four hundred and eighty-four records entered full-text eligibility "
               "triage. In the priority human pass, 19 records had open full text verified, "
               "16 had abstract-level verification, and 5 were not retrieved. Quantitative "
               "effect information was identified for 28 priority records. Exercise ranked "
               "highest for human healthspan-oriented evidence; other domains showed signals "
               "requiring further verification.")
    bold_label(doc, "Conclusion: ",
               "The extracted evidence supports healthspan-oriented exercise most strongly, "
               "while pharmacologic, senolytic, NAD+/sirtuin, microbiome, fasting, caloric "
               "restriction, and regenerative claims remain less certain. No intervention in "
               "this evidence set proves human age reversal. Final submission requires manual "
               "full-text eligibility confirmation, completed risk-of-bias assessment, and "
               "adjudication of duplicate records.")
    doc.add_paragraph()
    bold_label(doc, "Keywords: ",
               "ageing; healthspan; rejuvenation; biological age; senolytics; rapamycin; "
               "systematic review; evidence map")
    doc.add_paragraph()

    # ── Introduction ───────────────────────────────────────────────
    heading(doc, "Introduction", 1)
    para(doc,
         "Anti-ageing interventions are often discussed as though improved healthy ageing, "
         "delayed biological ageing, biomarker reversal, and rejuvenation were interchangeable. "
         "This creates a risk of clinical overclaiming. A rigorous review must separate human "
         "functional outcomes from animal lifespan results, cellular mechanisms, and surrogate "
         "biomarker shifts.[1,2]")
    para(doc,
         "Human studies in the extracted evidence set include multidomain lifestyle, resistance "
         "or multimodal exercise, microbiome modulation, exergame-based frailty prevention, "
         "prospective physical-activity cohorts, fasting-mimicking diet, topical rapamycin, and "
         "rapamycin healthspan records.[3,4,5,6,7,8,9,10]")
    para(doc,
         "This manuscript presents a reproducible evidence map and credibility ranking of "
         "anti-ageing and age-reversal intervention domains. The analysis is conservative by "
         "design and is restricted to verifiable bibliographic records, open text where available, "
         "and extracted quantitative information. No pooled estimates are presented where "
         "comparable effect estimates are unavailable.")
    para(doc,
         "The review question was not whether ageing has been definitively reversed in humans. "
         "It was whether any intervention class has credible evidence for healthspan improvement, "
         "lifespan extension, slowing of biological-age markers, biomarker reversal, or plausible "
         "rejuvenation. Those categories are treated as distinct evidentiary claims.")

    # ── Methods ────────────────────────────────────────────────────
    heading(doc, "Methods", 1)
    para(doc,
         "A PRISMA-style pipeline searched PubMed, Europe PMC, and Crossref; deduplicated "
         "records by DOI, PMID, and normalised title; screened title/abstract metadata; "
         "classified intervention and outcome domains; assessed model system and mechanism; "
         "and generated evidence credibility rankings. The study protocol was not prospectively "
         "registered because this is a pilot evidence-mapping exercise; prospective registration "
         "is recommended before a full-scale rerun. The PRISMA-style flow is shown in Figure 1.")
    para(doc,
         "Human evidence was prioritised for verification. Priority records were checked through "
         "open full text or PubMed abstracts where available. Risk-of-bias domains were assessed "
         "in a structured preliminary format. Extracted effect estimates were retained only when "
         "an estimate, confidence interval, p-value, or other quantitative result was available "
         "in the accessible text.")
    para(doc,
         "Meta-analysis was not performed because final verified effect estimates, uncertainty "
         "measures, denominators, intervention dose/duration, comparator details, and harmonised "
         "outcome definitions are not yet complete.")
    para(doc,
         "Intervention credibility combined number of extracted records, human evidence, human "
         "trial evidence, direct ageing or healthspan outcomes, biomarker evidence, surrogate "
         "burden, and hype-language burden. The scoring ranks evidence credibility; it does not "
         "create treatment recommendations. Additional records identified in the search but not "
         "individually discussed in the main text are listed in Supplementary Table S1.[15-40]")

    # ── Results ────────────────────────────────────────────────────
    heading(doc, "Results", 1)
    para(doc,
         "The pilot search retrieved 1155 raw records and produced 1029 deduplicated records. "
         "Title/abstract screening classified 29 records as include, 455 as uncertain, and 545 "
         "as exclude. Four hundred and eighty-four records entered full-text eligibility triage. "
         "Priority human verification was attempted for 40 records: 19 had open full text "
         "available, 16 had PubMed abstract-level verification, and 5 could not be verified "
         "through the automated open-source workflow.")
    para(doc,
         "The credibility ranking is shown in Table 1 and Figure 2. Exercise ranked first and "
         "was the only intervention domain categorised as the strongest current human healthspan "
         "signal (credibility score 31.48; 37 human records; hype rate 0.01). Microbiome "
         "(21.50), rapamycin/mTOR (21.24), senolytics (19.94), caloric restriction (19.52), "
         "lifestyle bundles (17.00), NAD+/sirtuin interventions (16.64), and fasting (15.10) "
         "showed human signals requiring verification. Supplements, metformin, sleep/circadian, "
         "controversial rejuvenation, reprogramming, and stem-cell approaches had lower scores "
         "or higher hype burden.")
    para(doc,
         "Representative priority human evidence with extracted effect estimates is summarised "
         "in Table 2. The records include frailty or function-oriented exercise and multidomain "
         "lifestyle interventions, microbiome modulation with cognitive outcomes, fasting-mimicking "
         "diet with biological-age markers, topical and systemic rapamycin records, caloric "
         "restriction, multidomain epigenetic-age intervention, and senolytic trial-design "
         "evidence.[3,4,5,6,7,8,9,10,11,12,13,14]")
    para(doc,
         "The clearest human functional signal came from physical-activity interventions. In a "
         "multimodal exercise trial in frail people with multiple sclerosis, the frailty index "
         "improved by −0.07 (95% CI −0.14 to −0.00) and mental-health quality of life improved "
         "by +21.24 points (95% CI 7.32 to 35.16).[2] An exergame-based exercise trial reported "
         "significant group-by-time improvements in handgrip strength and frailty outcomes.[4] "
         "A prospective five-cohort analysis suggested that physical activity modified the "
         "frailty–mortality association; hazard ratios were not re-extracted in the present "
         "synthesis.[6] These results support physical activity as healthspan-relevant evidence, "
         "not as proof of biological age reversal.")
    para(doc,
         "Dietary and microbiome interventions showed more mixed results. The PROMOTe randomised "
         "trial found no significant effect on the primary outcome of chair-rise time "
         "(β=0.579; 95% CI −1.08 to 2.24; p=0.494), but a favourable cognitive factor estimate "
         "(β=−0.482; 95% CI −0.813 to −0.141; p=0.014).[3] Fasting-mimicking diet was "
         "associated with a median biological-age reduction of 2.5 years after three cycles; "
         "this is a biomarker result and should be interpreted separately from clinical "
         "rejuvenation.[7] The caloric-restriction pilot record contributed primarily feasibility "
         "and adherence information.[12] A multidomain lifestyle intervention reported significant "
         "improvements in functional trajectories and epigenetic ageing markers.[13]")
    para(doc,
         "Rapamycin/mTOR evidence was heterogeneous. Topical rapamycin reduced p16INK4A protein "
         "levels (p=0.008) and increased collagen VII in human skin.[8] The PEARL trial reported "
         "improvements in lean tissue mass, pain, and general health; visceral adiposity was "
         "unchanged (ηp²=0.001; p=0.942).[9] Senolytic evidence in the priority table was "
         "represented by a trial-design record without outcome effect data.[14]")
    para(doc,
         "Preliminary risk-of-bias prompts are summarised in Table 3. These are not final "
         "RoB 2 or ROBINS-I judgements. The hype-versus-evidence comparison is shown in "
         "Figure 3. The translational readiness matrix is shown in Figure 4.")

    # ── TABLE 1 ───────────────────────────────────────────────────
    doc.add_paragraph()
    build_table1(doc)

    # ── FIGURE 1: PRISMA flow ─────────────────────────────────────
    doc.add_paragraph()
    embed_figure(
        doc,
        path=f"{FIGURES}/prisma_flow.png",
        width_inches=5.2,
        caption=(
            "Figure 1. PRISMA-style pilot flow diagram. Records through retrieval (1155 raw), "
            "deduplication (1029), title/abstract screening (29 include; 455 uncertain; "
            "545 exclude), and full-text eligibility triage (484 records). Counts reflect "
            "the capped pilot run; final counts will change after full-scale search rerun "
            "and manual adjudication."
        ),
    )

    # ── TABLE 2 ───────────────────────────────────────────────────
    doc.add_paragraph()
    build_table2(doc)

    # ── FIGURE 2: Evidence score ranking ──────────────────────────
    doc.add_paragraph()
    embed_figure(
        doc,
        path=f"{FIGURES}/evidence_score_ranking.png",
        width_inches=6.0,
        caption=(
            "Figure 2. Intervention credibility score ranking (n=14 classes). Scores "
            "incorporate human record count, human trial count, direct ageing/healthspan "
            "outcomes, biomarker count, surrogate burden, and hype-language penalty. "
            "Exercise ranked first (31.48); stem-cell approaches ranked last (1.68). "
            "Rankings are provisional pending final full-text adjudication."
        ),
    )

    # ── TABLE 3 ───────────────────────────────────────────────────
    doc.add_paragraph()
    build_table3(doc)

    # ── FIGURE 3: Hype vs evidence ────────────────────────────────
    doc.add_paragraph()
    embed_figure(
        doc,
        path=f"{FIGURES}/hype_vs_evidence_map.png",
        width_inches=6.0,
        caption=(
            "Figure 3. Hype-versus-evidence map. Credibility score (x-axis) vs "
            "hype-language burden (y-axis; proportion of extracted records flagged for "
            "hype-heavy terminology) by intervention domain. Bubble size = number of "
            "extracted records. Interventions with high credibility and low hype (exercise, "
            "lifestyle bundle) appear lower-right; reprogramming and controversial "
            "approaches appear upper-left."
        ),
    )

    # ── FIGURE 4: Translational readiness ────────────────────────
    doc.add_paragraph()
    embed_figure(
        doc,
        path=f"{FIGURES}/translational_matrix.png",
        width_inches=6.0,
        caption=(
            "Figure 4. Translational readiness matrix. Conservative category assignment "
            "for 14 intervention classes: A – healthspan support signal; "
            "B – promising, not recommendation-ready; C – biomarker/indirect signal; "
            "D – speculative/low directness. Categories are provisional pending "
            "final full-text adjudication."
        ),
    )

    # ── Discussion ────────────────────────────────────────────────
    heading(doc, "Discussion", 1)
    para(doc,
         "The main conclusion is deliberately conservative. In the extracted evidence, exercise "
         "has the clearest human healthspan signal, supporting functional healthy-ageing benefit "
         "rather than age reversal.[2,4,5,6]")
    para(doc,
         "Several other interventions have plausible biological or early human signals, but the "
         "present evidence does not justify clinical rejuvenation claims. Rapamycin/mTOR, "
         "senolytics, NAD+/sirtuin approaches, microbiome modulation, caloric restriction, and "
         "fasting should be described as promising or hypothesis-supporting until final "
         "full-text eligibility, risk of bias, effect estimates, dosing, comparator details, "
         "and safety outcomes are manually confirmed.[3,7,8,9,12,13,14]")
    para(doc,
         "The evidence map illustrates why clinical translation should remain conservative. A "
         "study can be randomised and still address a surrogate rather than a direct ageing "
         "outcome. A biomarker may shift in a favourable direction without demonstrating durable "
         "functional rejuvenation. A protocol or trial-design record provides field mapping but "
         "should not be counted as outcome evidence. Table 1 ranks intervention domains; Table 2 "
         "summarises representative effect estimates; Table 3 shows preliminary risk-of-bias "
         "prompts.")
    para(doc,
         "The effect-estimate table is intentionally not a meta-analysis table. It includes "
         "heterogeneous estimates — confidence intervals for frailty or cognitive outcomes, "
         "p-values for biomarker changes, feasibility information, and trial-design records with "
         "no outcome effect. Combining these outcomes would imply comparability that the "
         "extracted data do not support. Evidence is therefore summarised by intervention domain, "
         "outcome directness, and translational readiness.")
    para(doc,
         "For clinical readers, the most important distinction is between healthspan evidence "
         "and age-reversal evidence. Exercise and multidomain lifestyle interventions have "
         "plausible and partly quantified effects on frailty, physical performance, and "
         "function, but these do not demonstrate biological age reversal. Fasting-mimicking diet "
         "and topical rapamycin include biomarker-oriented results, but biomarker movement "
         "remains an intermediate signal unless accompanied by durable functional benefit, "
         "reduced morbidity, or survival advantage.")
    para(doc,
         "For senolytics, reprogramming, stem-cell/regenerative approaches, plasma-based "
         "interventions, and NAD+/sirtuin supplementation, mechanistic plausibility is not the "
         "same as clinical readiness. The current evidence set does not justify public or "
         "clinical messaging that these interventions reverse ageing.")
    para(doc,
         "A more definitive version of this review should extract denominators, baseline and "
         "follow-up means, standard deviations, between-group contrasts, adverse events, dose, "
         "comparator intensity, and follow-up duration for each eligible full text, enabling "
         "domain-specific meta-analysis for comparable outcomes such as grip strength, gait "
         "speed, frailty scores, epigenetic-age change, or selected biomarker endpoints.")
    para(doc,
         "Healthy-ageing recommendations should remain grounded in interventions with "
         "reproducible human functional evidence, especially physical activity and multidomain "
         "lifestyle support. Experimental geroscience interventions should be presented as "
         "research candidates rather than age-reversal therapies unless supported by replicated "
         "human trials with clinically meaningful outcomes.")

    # ── Strengths and Limitations ──────────────────────────────────
    heading(doc, "Strengths and Limitations", 1)
    para(doc,
         "Strengths include a reproducible workflow, explicit claim separation, conservative "
         "scoring, human-evidence prioritisation, full-text eligibility triage, duplicate-cohort "
         "checks, preliminary risk-of-bias prompts, and figures designed to separate hype from "
         "evidence.")
    para(doc,
         "Limitations include capped search retrieval preventing PRISMA-complete coverage; "
         "incomplete manual full-text review (484 records require final human adjudication); "
         "preliminary rather than final risk-of-bias judgements; absence of quantitative "
         "meta-analysis; 84 unresolved duplicate-title or overlapping-cohort groups requiring "
         "manual review; absence of prospective protocol registration; and PubMed and Crossref "
         "abstract gaps that reduce screening certainty.")

    # ── Conclusion ────────────────────────────────────────────────
    heading(doc, "Conclusion", 1)
    para(doc,
         "The extracted evidence supports a cautious hierarchy: exercise has the strongest "
         "human healthspan signal, while pharmacologic, nutritional, senolytic, NAD+/sirtuin, "
         "microbiome, and regenerative approaches remain promising but unproven for clinical "
         "age reversal. Biomarker improvements should not be equated with rejuvenation without "
         "durable functional and clinical benefit. Final conclusions await manual full-text "
         "verification, completed risk-of-bias assessment, and adjudicated duplicate-cohort "
         "review.")

    # ── Acknowledgements / Declarations ──────────────────────────
    heading(doc, "Acknowledgements", 1)
    para(doc, "None declared.")

    heading(doc, "Ethics", 1)
    para(doc,
         "Ethics committee approval was not required because this review used public "
         "bibliographic metadata and open-access text only, with no individual participant data.")

    heading(doc, "Funding", 1)
    para(doc, "No external funding is declared.")

    heading(doc, "Conflicts of Interest", 1)
    para(doc, "None declared.")

    heading(doc, "Data Availability", 1)
    para(doc,
         "All pipeline scripts, configuration files, results tables, figures, screening "
         "artifacts, risk-of-bias prompts, and submission-package documents are openly "
         "available at the project repository: "
         + GITHUB_URL +
         ". Additional records identified in the search are listed in Supplementary Table S1. "
         "Raw full-text cache files (PMC XML, PubMed XML) are not redistributed owing to "
         "publisher licence constraints; they can be re-downloaded by running the pipeline "
         "retrieval phase as documented in the repository README.")

    # ── References ────────────────────────────────────────────────
    heading(doc, "References", 1)
    refs = [
        "1. Yap KH, Chen C, Chong EJY, Kandiah N, Kivipelto M, Lai MKP, Maier AB, Matchar DB, Phua AKS, Teo CKL, Xu X, Zhou HJ, Chen CPL. Baseline characteristics of the SINGER multidomain dementia prevention randomized controlled trial and insights from the recruitment process. Alzheimers Dement. 2026. doi:10.1002/alz.71313.",
        "2. Zanotto T, Tabatabaei A, Lynch SG, He J, Lysaught M, Ahmadnezhad P, Ibude J, Devos H, Chaves LD, Troen BR, Sosnoff JJ. Reducing frailty in frail people with multiple sclerosis: Feasibility of a 6-week multimodal exercise training program. PLoS One. 2026. doi:10.1371/journal.pone.0347063.",
        "3. Ni Lochlainn M, Bowyer RCE, Moll JM, Garcia MP, Wadge S, Baleanu AF, Nessa A, Sheedy A. Effect of gut microbiome modulation on muscle function and cognition: the PROMOTe randomised controlled trial. Nat Commun. 2024. PMID:38424099.",
        "4. Tuan SH, Chang LH, Sun SF, Li CH, Chen GB, Tsai YJ. Assessing the Clinical Effectiveness of an Exergame-Based Exercise Training Program Using Ring Fit Adventure to Prevent and Postpone Frailty and Sarcopenia Among Older Adults in Rural Long-Term Care Facilities: Randomized Controlled Trial. J Med Internet Res. 2024. PMID:39024000.",
        "5. Yoon DH, Lee JY, Song W. Effects of Resistance Exercise Training on Cognitive Function and Physical Performance in Cognitive Frailty: A Randomized Controlled Trial. J Nutr Health Aging. 2018. doi:10.1007/s12603-018-1090-9.",
        "6. Zhu Z, Zhou X, Chen M, Dou C, Liu D, Kong L, Ye C, Xu M, Xu Y, Li M, Zhao Z, Zheng J, Lu J, Chen Y, Wang W, Ning G, Bi Y, Wang T. Interaction between physical activity and deficit-based frailty on all-cause mortality in older adults: a prospective study of five population-based cohorts. Lancet Reg Health West Pac. 2026. doi:10.1016/j.lanwpc.2025.101780.",
        "7. Brandhorst S, Levine ME, Wei M, Shelehchi M, Morgan TE, Nayak KS, Dorff T, Hong K. Fasting-mimicking diet causes hepatic and blood markers changes indicating reduced biological age and disease risk. Nat Commun. 2024. PMID:38378685.",
        "8. Chung CL, Lawrence I, Hoffman M, Elgindi D, Nadhan K, Potnis M, Jin A, Sershon C. Topical rapamycin reduces markers of senescence and aging in human skin: an exploratory, prospective, randomized trial. GeroScience. 2019. PMID:31761958.",
        "9. Moel M, Harinath G, Lee V, Nyquist A, Morgan SL, Isman A, Zalzala S. Influence of rapamycin on safety and healthspan metrics after one year: PEARL trial results. Aging. 2025. PMID:40188830.",
        "10. Brandhorst S, Choi IY, Wei M, Cheng CW, Sedrakyan S, Navarrete G, Dubeau L, Yap LP. A Periodic Diet that Mimics Fasting Promotes Multi-System Regeneration, Enhanced Cognitive Performance, and Healthspan. Cell Metab. 2015. PMID:26094889.",
        "11. Collins J, Longhurst G, Roschel H, Gualano B. Resistance training and co-supplementation with creatine and protein in older subjects with frailty. J Frailty Aging. 2016. doi:10.14283/jfa.2016.85.",
        "12. Houston DK, Fanning J, Nicklas BJ, Delany JP, Hsu FC, Chen SH, Walkup M, Neiberg RH, Stowe CL, Kennedy K, Espeland MA, Ard JD, Miller ME, Rejeski WJ, Kritchevsky SB. Caloric restriction and time-restricted eating in older adults with overweight or obesity: The HALLO Pilot Study. J Gerontol A Biol Sci Med Sci. 2026. doi:10.1093/gerona/glag061.",
        "13. Olaso-Gonzalez G, Millan-Domingo F, Garcia-Fernandez L, Garcia-Tercero E, Cebrian M, Garcia-Dominguez C, Carbonell JA, Casabo-Valles G, Garcia-Gimenez JL, Tamayo-Torres E, Gambini J, Tarazona-Santabalbina FJ, Vina J, Gomez-Cabrera MC. A Multidomain Lifestyle Intervention Is Associated With Improved Functional Trajectories and Favorable Changes in Epigenetic Aging Markers in Frail Older Adults: A Randomized Controlled Trial. Aging Cell. 2026. doi:10.1111/acel.70376.",
        "14. Ji J, Crespi CM, Yee L, Zekster YA, Al-Saleem A, Petersen L, Lee C, Son N, Smith C, Evans T, Tchkonia T, Kirkland JL, Kuchel GA, Cohen HJ, Sedrak MS. A phase II randomized placebo-controlled study of fisetin to improve physical function in breast cancer survivors: the TROFFi study rationale and trial design. Ther Adv Med Oncol. 2026. doi:10.1177/17588359261424668.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(10)

    p = doc.add_paragraph()
    r = p.add_run(
        "References 15–40: Additional records identified in the search are listed in "
        "Supplementary Table S1 and cited collectively in the Methods section."
    )
    r.italic = True
    r.font.size = Pt(10)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build_manuscript()
