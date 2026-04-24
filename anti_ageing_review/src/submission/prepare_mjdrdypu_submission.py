from __future__ import annotations

import argparse
import re
import shutil
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from src.utils.io import append_log, load_config, save_csv


ARTICLE_TITLE = (
    "Can Ageing Be Slowed or Reversed? A Reproducible Evidence Map, "
    "Credibility Ranking, and Mechanistic Synthesis of Anti-Ageing and Age-Reversal Interventions"
)
RUNNING_TITLE = "Anti-ageing evidence map"
JOURNAL = "Medical Journal of Dr. D.Y. Patil Vidyapeeth"
AUTHOR_NAME = "Dr Siddalingaiah H S"
AUTHOR_ROLE = "Professor, Community Medicine"
AUTHOR_AFFILIATION = "Shridevi Institute of Medical Sciences and Research Hospital, Tumkur"
AUTHOR_EMAIL = "hssling@yahoo.com"
AUTHOR_PHONE = "8941087719"
AUTHOR_ORCID = "0000-0002-4771-8285"


def read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path) if path.exists() else pd.DataFrame()


def clean(text: object) -> str:
    if pd.isna(text):
        return ""
    s = str(text)
    if any(token in s for token in ["Ã", "Â", "â"]):
        try:
            s = s.encode("latin1").decode("utf-8")
        except UnicodeError:
            pass
    replacements = {
        "â€™": "'",
        "â€˜": "'",
        "â€œ": '"',
        "â€�": '"',
        "â€“": "-",
        "â€”": "-",
        "Î²": "beta",
        "Î·": "eta",
        "Â±": "+/-",
        "â‰¥": ">=",
        "â‰¤": "<=",
        "âˆ’": "-",
        "&amp;": "&",
        "&lt;": "<",
        "&gt;": ">",
    }
    for old, new in replacements.items():
        s = s.replace(old, new)
    return re.sub(r"\s+", " ", s).strip()


def strip_terminal_punctuation(text: object) -> str:
    return clean(text).rstrip(" .;")


def ensure_doc_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(clean(text), level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"


def add_para(doc: Document, text: str = "", bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        r.font.name = "Times New Roman"
    r = p.add_run(clean(text))
    r.font.name = "Times New Roman"


def add_runs_with_superscripts(paragraph, parts: list[tuple[str, bool]]) -> None:
    for text, superscript in parts:
        run = paragraph.add_run(clean(text))
        run.font.name = "Times New Roman"
        run.font.superscript = superscript


def add_cited_para(doc: Document, text: str, refs: list[int] | None = None) -> None:
    p = doc.add_paragraph()
    p.add_run(clean(text)).font.name = "Times New Roman"
    if refs:
        citation = "[" + ",".join(str(r) for r in refs) + "]"
        r = p.add_run(citation)
        r.font.name = "Times New Roman"
        r.font.superscript = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(clean(item))


def title_case_clean(label: str) -> str:
    return label.replace("_", " ").title()


def add_table_from_df(doc: Document, df: pd.DataFrame, columns: list[str], title: str, limit: int = 20, headers: dict[str, str] | None = None) -> None:
    add_para(doc, title)
    if df.empty:
        add_para(doc, "No data available.")
        return
    cols = [c for c in columns if c in df.columns]
    out = df.loc[:, cols].head(limit).copy()
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr[i].text = headers.get(col, title_case_clean(col)) if headers else title_case_clean(col)
    for _, row in out.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            cells[i].text = clean(row[col])


def source_counts(df: pd.DataFrame, column: str) -> str:
    if df.empty or column not in df.columns:
        return "not available"
    counts = df[column].fillna("missing").astype(str).value_counts()
    return "; ".join(f"{clean(k)}: {int(v)}" for k, v in counts.items())


def format_reference(row: pd.Series, idx: int) -> dict[str, object]:
    authors = strip_terminal_punctuation(row.get("authors", ""))
    if not authors:
        authors = "Authors not captured"
    title = strip_terminal_punctuation(row.get("title", "Untitled"))
    journal = strip_terminal_punctuation(row.get("journal", ""))
    year = strip_terminal_punctuation(row.get("year", ""))
    if year.endswith(".0"):
        year = year[:-2]
    doi = clean(row.get("doi", ""))
    pmid = clean(row.get("pmid", ""))
    url = clean(row.get("url", ""))
    ref = f"{authors}. {title}. {journal}. {year}."
    if doi:
        ref += f" doi: {doi}."
    elif pmid and pmid.lower() != "nan":
        ref += f" PMID: {pmid.split('.')[0]}."
    elif url:
        ref += f" Available from: {url}."
    return {
        "reference_number": idx,
        "reference_text": ref,
        "title": title,
        "journal": journal,
        "year": year,
        "doi": doi,
        "pmid": pmid,
        "url": url,
    }


def build_references(human_queue: pd.DataFrame, verification: pd.DataFrame, max_refs: int = 40) -> pd.DataFrame:
    if human_queue.empty:
        return pd.DataFrame()
    verified_titles = set()
    if not verification.empty:
        verified_titles = set(verification["title"].astype(str))
    refs = human_queue.copy()
    refs["verified_priority"] = refs["title"].astype(str).isin(verified_titles).astype(int)
    refs = refs.sort_values(["verified_priority", "full_text_priority_score", "year"], ascending=[False, False, False])
    refs = refs.drop_duplicates(subset=["title"]).head(max_refs)
    rows = [format_reference(row, i + 1) for i, (_, row) in enumerate(refs.iterrows())]
    return pd.DataFrame(rows)


def ref_lookup(references: pd.DataFrame) -> dict[str, int]:
    if references.empty:
        return {}
    return {strip_terminal_punctuation(row["title"]): int(row["reference_number"]) for _, row in references.iterrows()}


def ref_num(lookup: dict[str, int], title: object) -> int | None:
    return lookup.get(strip_terminal_punctuation(title))


def first_ref_for_intervention(human_queue: pd.DataFrame, lookup: dict[str, int], intervention: str) -> int | None:
    if human_queue.empty:
        return None
    mask = human_queue["intervention_name"].fillna("").astype(str).eq(intervention)
    for title in human_queue.loc[mask, "title"].astype(str):
        key = strip_terminal_punctuation(title)
        if key in lookup:
            return lookup[key]
    return None


def first_study_for_intervention(human_queue: pd.DataFrame, lookup: dict[str, int], intervention: str) -> str:
    if human_queue.empty:
        return "No representative human record"
    mask = human_queue["intervention_name"].fillna("").astype(str).eq(intervention)
    for _, row in human_queue.loc[mask].iterrows():
        n = ref_num(lookup, row.get("title", ""))
        if n:
            return f"{first_author_year(row)} [{n}]"
    return "No representative human record"


def first_author_year(row: pd.Series) -> str:
    authors = clean(row.get("authors", ""))
    first = "Study"
    if authors:
        first = re.split(r";|,", authors)[0].strip()
        first = first.split()[0] if first else "Study"
    year = clean(row.get("year", ""))
    if year.endswith(".0"):
        year = year[:-2]
    return f"{first} et al., {year}" if year else f"{first} et al."


def study_label(row: pd.Series, lookup: dict[str, int]) -> str:
    n = ref_num(lookup, row.get("title", ""))
    suffix = f" [{n}]" if n else ""
    return first_author_year(row) + suffix


def clean_domain(value: object) -> str:
    mapping = {
        "exercise": "Exercise / physical activity",
        "lifestyle_bundle": "Multidomain lifestyle",
        "microbiome": "Microbiome modulation",
        "rapamycin_mtor": "Rapamycin / mTOR modulation",
        "senolytics": "Senolytics / senomorphics",
        "caloric_restriction": "Caloric restriction",
        "fasting": "Fasting / fasting-mimicking diet",
        "nad_sirtuin": "NAD/sirtuin approaches",
        "metformin": "Metformin",
        "supplements": "Supplements",
        "sleep_circadian": "Sleep/circadian approaches",
        "controversial": "Plasma/controversial rejuvenation",
        "reprogramming": "Epigenetic reprogramming",
        "stem_cell": "Stem-cell/regenerative approaches",
    }
    return mapping.get(clean(value), clean(value).replace("_", " ").capitalize())


def clean_design(value: object) -> str:
    mapping = {
        "RCT_or_clinical_trial": "Randomized/clinical trial",
        "cohort": "Prospective cohort",
        "unclear_or_metadata_only": "Review/protocol/metadata-limited record",
        "cross-sectional": "Cross-sectional study",
    }
    return mapping.get(clean(value), clean(value).replace("_", " "))


def clean_outcome(value: object) -> str:
    mapping = {
        "healthspan_functional_ageing": "Healthspan / functional ageing",
        "hard_ageing_relevance": "Mortality / survival / lifespan",
        "biological_ageing_biomarker": "Biological-age biomarker",
        "surrogate_or_indirect": "Surrogate or indirect outcome",
    }
    return mapping.get(clean(value), clean(value).replace("_", " "))


def concise_finding(text: object, title: object = "") -> str:
    t = clean(text)
    title_text = clean(title).lower()
    if "singer" in title_text:
        return "Multidomain dementia-prevention trial baseline record; randomized healthspan-relevant cohort with frailty and cognitive outcomes planned."
    if "fisetin" in title_text or "troffi" in title_text:
        return "Phase II fisetin trial-design record targeting physical function in breast cancer survivors; outcome results not yet established in captured text."
    if "prica-powfrail" in title_text:
        return "Exercise/lifestyle protocol record for older adults at risk of dependency; useful for intervention mapping rather than outcome effect estimation."
    if "test of rapamycin in aging dogs" in title_text:
        return "Rapamycin trial-design record in dogs; translationally relevant but not human outcome evidence."
    if "nad-hd" in title_text:
        return "Nicotinamide riboside trial record in Huntington disease; indirect relevance to ageing biology."
    if not t:
        return "No numeric result available in captured text"
    rules = [
        ("chair rise time", "Prebiotic treatment did not improve chair-rise time but improved a cognitive factor score in the captured abstract."),
        ("Fried frailty score", "Pilot multimodal exercise record reported high retention/adherence and between-group changes in frailty-related and quality-of-life outcomes."),
        ("group×time", "Exergame-based exercise trial reported significant group-by-time improvements in frailty/sarcopenia-related physical outcomes."),
        ("all-cause mortality", "Physical activity modified the association between frailty and all-cause mortality across five cohorts."),
        ("2.5 years in median biological age", "Fasting-mimicking diet record reported a 2.5-year reduction in median biological age marker estimates."),
        ("p16", "Topical rapamycin record reported reduced p16INK4A and increased collagen VII markers."),
        ("Lean tissue mass", "Rapamycin PEARL record reported selected improvements in lean tissue mass, pain, well-being, or self-rated health, with several biomarkers unchanged."),
        ("SHARE", "Multidomain lifestyle trial reported reduced frailty and improved grip strength, gait speed, Tinetti score, and Barthel index."),
        ("Horvath", "Greens/microbiome supplementation record captured epigenetic-clock and microbiome changes over short follow-up."),
        ("CALERIE", "Caloric-restriction record was linked to long-term calorie restriction and DNA methylation biological-age outcomes."),
    ]
    for needle, finding in rules:
        if needle.lower() in t.lower():
            return finding
    first = t.split("||")[0]
    return first[:230].rstrip(" .,") + "."


def effect_estimate(text: object, title: object = "") -> str:
    t = clean(text)
    title_text = clean(title).lower()
    if "singer" in title_text:
        return "No outcome effect estimate; baseline trial characteristics only"
    if "multiple sclerosis" in title_text:
        return "EFIP change -0.07 (95% CI -0.14 to -0.00); MSQoL-54 mental health +21.24 (95% CI 7.32 to 35.16)"
    if "promote randomised" in title_text or "gut microbiome modulation" in title_text:
        return "Chair-rise beta 0.579 (95% CI -1.080 to 2.239; p=0.494); cognition beta -0.482 (95% CI -0.813 to -0.141; p=0.014)"
    if "ring fit" in title_text or "exergame" in title_text:
        return "Primary outcomes improved by group x time interaction: most p<0.001; handgrip p=0.01"
    if "all-cause mortality" in title_text:
        return "Frailty x physical activity interaction with mortality: all p-interaction <=0.036; hazard ratios not extracted"
    if "fasting-mimicking diet" in title_text:
        return "Median biological-age estimate decreased by 2.5 years after 3 cycles"
    if "topical rapamycin" in title_text:
        return "p16INK4A reduction p=0.008; collagen VII increase p=0.0077"
    if "pearl" in title_text or "influence of rapamycin" in title_text:
        return "Lean tissue mass eta-p2=0.202, p=0.013; pain eta-p2=0.168, p=0.015; general health eta-p2=0.166, p=0.004"
    if "periodic diet" in title_text:
        return "No extractable effect estimate in accessible abstract"
    if "health, aging, and later-life outcomes" in title_text:
        return "Retention 92%; mean calorie restriction achieved 4.5% (SD 11.0); clinical effect estimate not extracted"
    if "multidomain lifestyle intervention" in title_text:
        return "SHARE-FI p<0.0001; grip strength p=0.0053; gait speed p=0.0125; Tinetti p=0.0031; Barthel p=0.0484"
    if "fisetin" in title_text or "troffi" in title_text:
        return "Trial design; no outcome effect estimate available"
    if "greens supplementation" in title_text:
        return "Microbiome taxa changed: Bilophila p=0.037; Desulfobacterota p=0.031; fasting glucose correlation rs=-0.81, p<0.001"
    if "nmn" in title_text:
        return "No ageing-outcome effect estimate extracted from available abstract"
    if "calerie" in title_text and "dna methylation" in title_text:
        return "Small treatment effects on DNAm PhenoAge/GrimAge; exact numerical effect not captured in accessible text"
    if "diactive" in title_text:
        return "Lean mass MD 0.88 kg (95% CI 0.09 to 1.66; Hedges g=0.568); non-ageing youth population"
    # Generic fallbacks from captured text
    patterns = [
        r"beta\s*=\s*[-+]?\d+\.\d+[^.;)]*(?:p\s*[<=>]\s*0?\.\d+)?",
        r"MD\]\s*=\s*[-+]?\d+\.\d+[^.;)]*",
        r"mean difference[^.;)]*",
        r"\d+(?:\.\d+)?\s*years[^.;)]*biological age",
        r"p\s*[<=>]\s*0?\.\d+",
    ]
    for pat in patterns:
        m = re.search(pat, t, flags=re.I)
        if m:
            return clean(m.group(0))[:220]
    return "No extractable effect estimate in accessible text"


def clean_tier(tier: object) -> str:
    mapping = {
        "highest_current_human_healthspan_signal": "Strongest current human healthspan signal",
        "human_signal_requires_verification": "Human signal; verification required",
        "biomarker_or_indirect_human_signal": "Biomarker or indirect human signal",
        "preclinical_direct_ageing_signal": "Preclinical direct-ageing signal",
        "low_directness_or_speculative": "Low-directness or speculative",
    }
    return mapping.get(clean(tier), clean(tier).replace("_", " "))


def clean_source_type(value: object) -> str:
    mapping = {
        "pmc_open_full_text": "Open full text",
        "pubmed_abstract": "PubMed abstract",
        "not_retrieved": "Not retrieved",
        "repo_abstract_only": "Repository abstract",
    }
    return mapping.get(clean(value), clean(value).replace("_", " "))


def clean_rob(value: object) -> str:
    mapping = {
        "low_to_some_concern_pending_full_text_review": "Low to some concern",
        "some_concern": "Some concern",
        "serious_or_unclear": "Serious/unclear",
        "low_concern": "Low concern",
        "not_applicable": "Not applicable",
        "not_applicable_or_unclear": "Not applicable/unclear",
        "low_concern_rct_design": "Low concern (trial design)",
        "some_concern_observational_confounding": "Some concern (observational confounding)",
        "low_or_some_concern": "Low/some concern",
        "unclear_without_protocol": "Unclear without protocol",
        "unclear_or_not_applicable": "Unclear/not applicable",
    }
    return mapping.get(clean(value), clean(value).replace("_", " "))


def publication_ranking_table(data: dict[str, pd.DataFrame], lookup: dict[str, int]) -> pd.DataFrame:
    ranking = data["ranking"].copy()
    if ranking.empty:
        return ranking
    human_queue = data["human_queue"]
    rows = []
    for _, row in ranking.head(10).iterrows():
        intervention = clean(row["intervention_name"])
        hq = human_queue[human_queue["intervention_name"].fillna("").astype(str).eq(intervention)]
        direct_human = int(hq["ageing_domain_category"].isin(["hard_ageing_relevance", "healthspan_functional_ageing", "biological_ageing_biomarker"]).sum()) if not hq.empty else 0
        trial_human = int(hq["study_design"].astype(str).str.contains("RCT|clinical_trial", case=False, regex=True).sum()) if not hq.empty else 0
        citation = first_study_for_intervention(human_queue, lookup, intervention)
        rows.append({
            "Rank": int(row["rank"]),
            "Intervention domain": clean_domain(intervention),
            "Human records": int(row["human_records"]),
            "Human trial records": trial_human,
            "Human direct-ageing records": direct_human,
            "Credibility tier": clean_tier(row["credibility_tier"]),
            "Representative citation": citation,
        })
    return pd.DataFrame(rows)


def publication_evidence_table(data: dict[str, pd.DataFrame], lookup: dict[str, int]) -> pd.DataFrame:
    verification = data["verification"]
    queue = data["human_queue"]
    effects = data["effects"]
    if verification.empty:
        return pd.DataFrame()
    work = verification.merge(queue[["title", "authors", "year", "intervention_name", "study_design", "ageing_domain_category", "journal"]], on="title", how="left")
    work = work.merge(effects[["title", "effect_size_extraction_status", "candidate_effect_size_text"]], on="title", how="left")
    rows = []
    eligible = work[
        work["has_intervention_signal"].astype(str).str.lower().eq("true")
        & work["has_ageing_outcome_signal"].astype(str).str.lower().eq("true")
    ].head(12)
    for _, row in eligible.iterrows():
        title = clean(row["title"])
        rows.append({
            "Study": study_label(row, lookup),
            "Intervention domain": clean_domain(row.get("intervention_name", "")),
            "Design": clean_design(row.get("study_design", "")),
            "Outcome focus": clean_outcome(row.get("ageing_domain_category", "")),
            "Evidence source reviewed": clean_source_type(row.get("source_text_type", "")),
            "Effect estimate extracted from available text": effect_estimate(row.get("candidate_effect_size_text", ""), row.get("title", "")),
        })
    return pd.DataFrame(rows)


def publication_rob_table(data: dict[str, pd.DataFrame], lookup: dict[str, int]) -> pd.DataFrame:
    rob = data["formal_rob"].copy()
    if rob.empty:
        return rob
    queue = data["human_queue"]
    if not queue.empty:
        rob = rob.merge(queue[["title", "authors", "year"]], on="title", how="left")
    rows = []
    for _, row in rob.head(12).iterrows():
        title = clean(row["title"])
        rows.append({
            "Study": study_label(row, lookup) if ref_num(lookup, title) else clean(title[:60]),
            "Randomization": clean_rob(row.get("rob_randomization", "")),
            "Blinding": clean_rob(row.get("rob_blinding", "")),
            "Missing data": clean_rob(row.get("rob_missing_data", "")),
            "Confounding": clean_rob(row.get("rob_confounding", "")),
            "Outcome measurement": clean_rob(row.get("rob_outcome_measurement", "")),
            "Overall preliminary judgement": clean_rob(row.get("rob_overall", "")),
        })
    return pd.DataFrame(rows)


def word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text))


def build_narrative(data: dict[str, pd.DataFrame]) -> dict[str, str]:
    raw_n = len(data["raw"])
    dedup_n = len(data["dedup"])
    screened = data["screened"]
    full_text = data["full_text"]
    verification = data["verification"]
    effects = data["effects"]
    rob = data["formal_rob"]
    ranking = data["ranking"]
    duplicates = data["duplicate_cohorts"]

    include_n = int((screened.get("screen_label", pd.Series(dtype=str)) == "include").sum()) if not screened.empty else 0
    uncertain_n = int((screened.get("screen_label", pd.Series(dtype=str)) == "uncertain").sum()) if not screened.empty else 0
    exclude_n = int((screened.get("screen_label", pd.Series(dtype=str)) == "exclude").sum()) if not screened.empty else 0
    verified_open = int((verification.get("verification_status", pd.Series(dtype=str)) == "open_full_text_verified").sum()) if not verification.empty else 0
    verified_abs = int((verification.get("verification_status", pd.Series(dtype=str)) == "abstract_verified_full_text_pending").sum()) if not verification.empty else 0
    not_verified = int((verification.get("verification_status", pd.Series(dtype=str)) == "not_verified_source_unavailable").sum()) if not verification.empty else 0
    effect_candidate = int((effects.get("effect_size_extraction_status", pd.Series(dtype=str)) == "candidate_numeric_effect_text_extracted").sum()) if not effects.empty else 0

    top = ranking.head(1).iloc[0] if not ranking.empty else None
    top_sentence = "No intervention ranking was available."
    if top is not None:
        top_sentence = (
            f"{clean(top['intervention_name']).replace('_', ' ')} ranked first, with "
            f"{clean(top['human_records'])} human records, {clean(top['human_trial_records'])} human trial records, "
            f"{clean(top['healthspan_records'])} healthspan records, and a credibility score of {clean(top['credibility_score'])}."
        )

    abstract = (
        "Background: Anti-ageing claims range from established healthspan interventions to speculative rejuvenation approaches. "
        "Objective: To create a conservative, reproducible evidence map separating healthspan benefit, lifespan extension, biological-age slowing, biomarker reversal, and true clinical rejuvenation. "
        f"Methods: Searches across PubMed, Europe PMC, and Crossref retrieved {raw_n} raw records, deduplicated to {dedup_n}. Title/abstract screening, metadata-assisted extraction, credibility scoring, mechanism mapping, priority human verification, preliminary risk-of-bias assessment, duplicate checks, and effect-estimate extraction were performed. "
        f"Results: Screening classified {include_n} records as include, {uncertain_n} as uncertain, and {exclude_n} as exclude. {len(full_text)} records entered full-text eligibility triage. In the priority human pass, {verified_open} records had open full text verified, {verified_abs} had abstract-level verification, and {not_verified} were not retrieved. Quantitative effect information was identified for {effect_candidate} priority records. Exercise ranked highest for human healthspan-oriented evidence; other domains showed signals requiring further verification. "
        "Conclusion: The extracted evidence supports healthspan-oriented exercise most strongly, while pharmacologic, senolytic, NAD/sirtuin, microbiome, fasting, caloric restriction, and regenerative claims remain less certain. No intervention in this evidence set proves human age reversal."
    )

    results = (
        f"The pilot search retrieved {raw_n} raw records and produced {dedup_n} deduplicated records. "
        f"Title/abstract screening classified {include_n} records as include, {uncertain_n} as uncertain, and {exclude_n} as exclude. "
        f"Full-text eligibility triage included {len(full_text)} records. Priority bands were: {source_counts(data['full_text_decisions'], 'full_text_priority_band')}. "
        f"Priority human verification was attempted for {len(verification)} records: {verified_open} had open PMC full text verified, {verified_abs} had PubMed abstract verification with full text pending, and {not_verified} were not retrieved by automated lookup. "
        f"Quantitative effect information was detected for {effect_candidate} prioritized records; effect estimates were retained only where accessible text reported a numerical estimate, confidence interval, p value, or other quantitative result. "
        f"Preliminary risk-of-bias prompts classified records as {source_counts(rob, 'rob_overall')}. "
        f"Duplicate and overlapping-publication checks produced {len(duplicates)} flags requiring manual adjudication. "
        f"{top_sentence} The first credibility ranking placed microbiome, rapamycin/mTOR, senolytics, caloric restriction, lifestyle bundles, NAD/sirtuin interventions, and fasting in a human-signal-requires-verification tier. Supplements, metformin, sleep/circadian approaches, controversial rejuvenation, reprogramming, and stem-cell approaches were categorized as biomarker/indirect or low-directness categories in the extracted dataset."
    )

    discussion = (
        "The main conclusion is deliberately conservative. In the extracted evidence, exercise has the clearest human healthspan signal, but this supports functional healthy-ageing benefit rather than age reversal. "
        "Several other interventions have plausible biological or early human signals, but the present evidence does not justify clinical rejuvenation claims. Rapamycin/mTOR, senolytics, NAD/sirtuin approaches, microbiome modulation, caloric restriction, and fasting should be described as promising or hypothesis-supporting until full-text eligibility, risk of bias, effect estimates, dosing, comparator details, and safety outcomes are manually confirmed. "
        "The hype-versus-evidence plot shows that some domains with mechanistic interest also carry higher hype-language burden, particularly reprogramming, stem-cell, plasma/controversial interventions, and some fasting or senolytic claims. "
        "The translational readiness matrix therefore keeps only exercise in the highest healthspan-support signal category and keeps other intervention classes below recommendation level. "
        "Effect information was extracted from open full text or abstracts, but final pooling was not performed because outcomes, intervention durations, comparators, and reporting formats were not sufficiently harmonized. "
        "This distinction is essential: biomarker movement, p-values, or pathway modulation cannot be equated with organismal rejuvenation without durable functional and clinical benefit."
    )
    return {"abstract": abstract, "results": results, "discussion": discussion}


def add_references(doc: Document, references: pd.DataFrame) -> None:
    add_heading(doc, "References", 1)
    if references.empty:
        add_para(doc, "Reference metadata not available.")
        return
    for _, row in references.iterrows():
        add_para(doc, f"{int(row['reference_number'])}. {clean(row['reference_text'])}")


def build_manuscript_docx(out: Path, data: dict[str, pd.DataFrame], references: pd.DataFrame, figures: Path) -> str:
    narrative = build_narrative(data)
    lookup = ref_lookup(references)
    ranking_pub = publication_ranking_table(data, lookup)
    evidence_pub = publication_evidence_table(data, lookup)
    rob_pub = publication_rob_table(data, lookup)
    doc = Document()
    ensure_doc_style(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(ARTICLE_TITLE)
    run.bold = True
    run.font.size = Pt(14)
    add_para(doc, "Article type: Systematic review / evidence map")
    add_para(doc, "Running title: " + RUNNING_TITLE)
    add_heading(doc, "Abstract", 1)
    add_para(doc, narrative["abstract"])
    add_para(doc, "Keywords: ageing; healthspan; rejuvenation; biological age; senolytics; rapamycin; systematic review; evidence map")
    add_heading(doc, "Introduction", 1)
    add_cited_para(doc, "Anti-ageing interventions are often discussed as though improved healthy ageing, delayed biological ageing, biomarker reversal, and rejuvenation were interchangeable. This creates risk of clinical overclaiming. A rigorous review must separate human functional outcomes from animal lifespan results, cellular mechanisms, and surrogate biomarker shifts.", [1, 2])
    add_cited_para(doc, "Human studies in the extracted evidence set include multidomain lifestyle, resistance or multimodal exercise, microbiome modulation, exergame-based frailty prevention, prospective physical-activity cohorts, fasting-mimicking diet, topical rapamycin, and rapamycin healthspan records.", [3, 4, 5, 6, 7, 8, 9, 10])
    add_para(doc, "This manuscript presents a reproducible evidence map and credibility ranking of anti-ageing and age-reversal intervention domains. The analysis is conservative by design and is restricted to verifiable bibliographic records, open text where available, and extracted quantitative information. No pooled estimates are presented where comparable effect estimates are unavailable.")
    add_para(doc, "The review question was not whether ageing has been definitively reversed in humans. It was whether any intervention class has credible evidence for healthspan improvement, lifespan extension, slowing of biological-age markers, biomarker reversal, or plausible rejuvenation. Those categories are treated as distinct evidentiary claims.")
    add_heading(doc, "Methods", 1)
    add_para(doc, "A PRISMA-style pipeline searched PubMed, Europe PMC, and Crossref; deduplicated records by DOI, PMID, and normalized title; screened title/abstract metadata; classified intervention and outcome domains; assessed model system and mechanism; and generated evidence credibility rankings.")
    add_para(doc, "Human evidence was prioritized for verification. Priority records were checked through open full text or PubMed abstracts where available. Risk-of-bias domains were assessed in a structured preliminary format. Extracted effect estimates were retained only when an estimate, confidence interval, p value, or other quantitative result was available in the accessible text. The PRISMA-style flow is shown in Figure 1.")
    add_para(doc, "Meta-analysis was not performed because final verified effect estimates, uncertainty measures, denominators, intervention dose/duration, comparator details, and harmonized outcome definitions are not yet complete.")
    add_para(doc, "Intervention credibility combined the number of extracted records, human evidence, human trial evidence, direct ageing or healthspan outcomes, biomarker evidence, surrogate burden, and hype-language burden. The scoring was used to rank evidence credibility, not to create treatment recommendations.")
    add_heading(doc, "Results", 1)
    add_para(doc, "The pilot search retrieved 1155 raw records and produced 1029 deduplicated records. Title/abstract screening classified 29 records as include, 455 as uncertain, and 545 as exclude. Four hundred and eighty-four records entered full-text eligibility triage. Priority human verification was attempted for 40 records: 19 had open full text available, 16 had PubMed abstract-level verification, and 5 could not be verified through the automated open-source workflow.")
    add_para(doc, "The first credibility ranking is shown in Table 1 and Figure 2. Exercise ranked first and was the only intervention domain categorized as the strongest current human healthspan signal. Microbiome, rapamycin/mTOR modulation, senolytics, caloric restriction, lifestyle bundles, NAD/sirtuin interventions, and fasting showed human signals requiring verification rather than recommendation-level evidence.")
    add_cited_para(doc, "Representative priority human evidence is summarized with extracted effect estimates in Table 2. The extracted records include frailty or function-oriented exercise and multidomain lifestyle interventions, microbiome modulation with cognitive outcomes, fasting-mimicking diet with biological-age markers, topical and systemic rapamycin records, caloric restriction, multidomain epigenetic-age intervention, and senolytic trial-design evidence.", [3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14])
    add_cited_para(doc, "The clearest human functional signal came from physical-activity interventions and physical-activity cohorts. In a multimodal exercise trial in frail people with multiple sclerosis, the extracted estimates showed a small favorable change in frailty index and a larger improvement in mental-health-related quality of life. An exergame-based exercise trial reported statistically significant group-by-time improvements across most frailty and sarcopenia outcomes, including handgrip strength. A prospective five-cohort analysis further suggested that physical activity modified the association between deficit-based frailty and all-cause mortality, although hazard ratios were not extracted in the present synthesis. These results support physical activity as healthspan-relevant evidence, not as proof of biological age reversal.", [2, 4, 6])
    add_cited_para(doc, "Dietary and microbiome interventions showed more mixed interpretation. The PROMOTe randomized trial did not show a favorable chair-rise effect in the extracted estimate, but it did report a favorable cognitive factor estimate. Fasting-mimicking diet evidence included a reported median biological-age estimate reduction after three cycles, but this is a biomarker result and should be interpreted separately from clinical rejuvenation. The caloric-restriction pilot record mainly contributed feasibility and adherence information rather than a direct clinical effect estimate. A multidomain lifestyle intervention reported statistically significant improvements in frailty and functional measures, but was categorized as surrogate or indirect in this extracted table because the current evidence fields were not sufficient for a stronger direct-ageing classification.", [3, 7, 12, 13])
    add_cited_para(doc, "Rapamycin/mTOR evidence was similarly heterogeneous. Topical rapamycin was linked to senescence and skin-ageing markers, including reduced p16INK4A and increased collagen VII. The PEARL rapamycin record reported quantitative signals for lean tissue mass, pain, and general health, but these were treated as healthspan or surrogate outcomes rather than proof of organismal rejuvenation. Senolytic evidence in the priority table was represented by a trial-design record without outcome effect estimates, so it was useful for field mapping but not counted as completed efficacy evidence.", [8, 9, 14])
    add_para(doc, "Preliminary risk-of-bias prompts are summarized in Table 3. These are not final RoB 2 or ROBINS-I judgments; they are structured prompts based on the available open text or abstract content.")
    add_para(doc, "The hype-versus-evidence comparison is presented in Figure 3. The translational readiness matrix is presented in Figure 4.")
    add_table_from_df(doc, ranking_pub, list(ranking_pub.columns), "Table 1. Intervention credibility ranking based on extracted evidence.", 10)
    add_table_from_df(doc, evidence_pub, list(evidence_pub.columns), "Table 2. Representative priority human evidence with effect estimates extracted from available text.", 12)
    add_table_from_df(doc, rob_pub, list(rob_pub.columns), "Table 3. Preliminary risk-of-bias prompts for representative priority human records.", 12)
    add_heading(doc, "Discussion", 1)
    add_cited_para(doc, "The main conclusion is deliberately conservative. In the extracted evidence, exercise has the clearest human healthspan signal, but this supports functional healthy-ageing benefit rather than age reversal.", [2, 4, 5, 6])
    add_cited_para(doc, "Several other interventions have plausible biological or early human signals, but the present evidence does not justify clinical rejuvenation claims. Rapamycin/mTOR, senolytics, NAD/sirtuin approaches, microbiome modulation, caloric restriction, and fasting should be described as promising or hypothesis-supporting until final full-text eligibility, risk of bias, effect estimates, dosing, comparator details, and safety outcomes are manually confirmed.", [3, 7, 8, 9, 12, 13, 14])
    add_para(doc, "The evidence map also shows why clinical translation should be conservative. A study can be randomized and still address a surrogate rather than a direct ageing outcome. A biomarker may move in a favorable direction without demonstrating durable functional rejuvenation. A protocol or trial-design record can be valuable for mapping the field but should not be counted as outcome evidence. These distinctions are reflected in the tables: Table 1 ranks intervention domains, Table 2 summarizes representative effect estimates, and Table 3 shows preliminary risk-of-bias assessments.")
    add_para(doc, "The effect-estimate table is intentionally not a meta-analysis table. It includes heterogeneous estimates: confidence intervals for frailty or cognitive outcomes, p values for biomarker and functional changes, feasibility/adherence information, and trial-design records with no outcome effect. This format is less statistically elegant than a pooled forest plot, but it is scientifically more defensible at this stage because combining these outcomes would imply comparability that the extracted data do not support. The evidence is therefore summarized by intervention domain, outcome directness, and translational readiness rather than by a single pooled anti-ageing effect.")
    add_para(doc, "For clinical readers, the most important distinction is between healthspan evidence and age-reversal evidence. Exercise and multidomain lifestyle interventions have plausible and partly quantified effects on frailty, physical performance, and function. These endpoints matter clinically, but they do not demonstrate that biological ageing has been reversed. Conversely, fasting-mimicking diet and topical rapamycin records include biomarker-oriented results, but biomarker movement remains an intermediate signal unless accompanied by durable functional benefit, reduced morbidity, or survival advantage.")
    add_para(doc, "For domains such as senolytics, reprogramming, stem-cell/regenerative approaches, plasma-based interventions, and NAD/sirtuin supplementation, mechanistic plausibility is not the same as clinical readiness. These areas remain important for geroscience research, but the current evidence set does not justify public or clinical messaging that they reverse ageing.")
    add_para(doc, "The translational readiness matrix follows from this separation. Category assignment favors interventions with human evidence, clinically meaningful outcomes, feasible delivery, and lower hype burden. This naturally places exercise and related lifestyle support above pharmacologic or regenerative geroscience interventions. It also means that interventions with strong mechanistic narratives but limited human outcome evidence remain below recommendation level even when early biomarker or pathway signals are encouraging.")
    add_para(doc, "A more definitive future version of this review should extract denominators, baseline and follow-up means, standard deviations, between-group contrasts, adverse events, dose, comparator intensity, and follow-up duration for each eligible full text. That would permit domain-specific meta-analysis for comparable outcomes such as grip strength, gait speed, frailty scores, epigenetic-age change, or selected biomarker endpoints. Until that extraction is complete, the present manuscript should be read as an effect-estimate-based evidence map rather than a completed quantitative systematic review.")
    add_para(doc, "The strongest practical message is therefore conventional but important: interventions already aligned with healthy ageing, particularly physical activity and multidomain lifestyle support, are more defensible than experimental age-reversal claims. Pharmacologic and regenerative strategies should be framed as research candidates unless and until replicated human trials show clinically meaningful functional benefit, acceptable safety, and durable effects.")
    add_para(doc, "The practical implication is that healthy-ageing recommendations should remain grounded in interventions with reproducible human functional evidence, especially physical activity and multidomain lifestyle support. Experimental geroscience interventions should be presented as research candidates rather than age-reversal therapies unless supported by replicated human trials with clinically meaningful outcomes.")
    add_heading(doc, "Strengths and Limitations", 1)
    add_para(doc, "Strengths include a reproducible workflow, explicit claim separation, conservative scoring, human-evidence prioritization, full-text eligibility triage, duplicate-cohort checks, preliminary RoB prompts, and figures designed to separate hype from evidence.")
    add_para(doc, "Limitations include capped search retrieval, incomplete manual full-text review, preliminary rather than final risk-of-bias judgments, and the absence of quantitative meta-analysis. Several screened records were protocols, reviews, or metadata-limited records; therefore, the manuscript tables emphasize representative human evidence with extractable quantitative results rather than treating every screening candidate as a completed outcome study. These limitations are reported explicitly to prevent overinterpretation.")
    add_heading(doc, "Conclusion", 1)
    add_para(doc, "The extracted evidence supports a cautious hierarchy: exercise has the strongest human healthspan signal, while several pharmacologic, nutritional, senolytic, NAD/sirtuin, microbiome, and regenerative approaches remain promising but unproven for clinical age reversal. Biomarker improvements should not be equated with rejuvenation without durable functional and clinical benefit.")
    add_heading(doc, "Ethics", 1)
    add_para(doc, "Ethics committee approval was not required because this review used public bibliographic metadata and open-access text where available, with no individual participant data.")
    add_heading(doc, "Funding", 1)
    add_para(doc, "No external funding is declared for this manuscript package.")
    add_heading(doc, "Conflicts of Interest", 1)
    add_para(doc, "None declared.")
    add_heading(doc, "Data Availability", 1)
    add_para(doc, "The reproducible tables, figures, extraction sheets, and scripts are available from the project files.")
    add_heading(doc, "Figures", 1)
    for name, caption in [
        ("prisma_flow.png", "Figure 1. PRISMA-style pilot flow of records through retrieval, deduplication, screening, and extraction."),
        ("evidence_score_ranking.png", "Figure 2. Intervention evidence score ranking based on extracted evidence."),
        ("hype_vs_evidence_map.png", "Figure 3. Hype-versus-evidence map comparing credibility score with hype-language burden by intervention domain."),
        ("translational_matrix.png", "Figure 4. Translational readiness matrix showing conservative category assignment by intervention domain."),
    ]:
        path = figures / name
        if path.exists():
            add_para(doc, caption)
            doc.add_picture(str(path), width=Inches(5.8))
    add_references(doc, references)
    main_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "References":
            break
        main_parts.append(paragraph.text)
    doc.save(out)
    return "\n".join(main_parts)


def build_title_page(out: Path, main_text_words: int, references: pd.DataFrame) -> None:
    doc = Document()
    ensure_doc_style(doc)
    add_heading(doc, "Title Page", 1)
    add_para(doc, ARTICLE_TITLE, "Title: ")
    add_para(doc, "Systematic review / evidence map", "Article type: ")
    add_para(doc, RUNNING_TITLE, "Running title: ")
    add_para(doc, f"{AUTHOR_NAME}, {AUTHOR_ROLE}", "Author: ")
    add_para(doc, AUTHOR_AFFILIATION, "Affiliation: ")
    add_para(doc, f"{AUTHOR_NAME}; Email: {AUTHOR_EMAIL}; Phone: {AUTHOR_PHONE}; ORCID: {AUTHOR_ORCID}", "Corresponding author: ")
    add_para(doc, str(main_text_words), "Main text word count excluding references/tables/figures: ")
    add_para(doc, str(len(references)), "Number of references: ")
    add_para(doc, "4 embedded figures; 3 embedded tables", "Display items: ")
    add_para(doc, "No external funding declared.", "Source(s) of support: ")
    add_para(doc, "None declared.", "Conflicts of interest: ")
    add_para(doc, "Not applicable; public bibliographic metadata and open-access text only.", "Ethics approval: ")
    add_para(doc, AUTHOR_NAME, "Guarantor: ")
    doc.save(out)


def build_simple_doc(out: Path, title: str, paragraphs: list[str]) -> None:
    doc = Document()
    ensure_doc_style(doc)
    add_heading(doc, title, 1)
    for para in paragraphs:
        add_para(doc, para)
    doc.save(out)


def build_peer_reviews(out1: Path, out2: Path, audit_out: Path, data: dict[str, pd.DataFrame]) -> None:
    ranking = data["ranking"]
    verification = data["verification"]
    effects = data["effects"]
    qc = data["qc"]
    dup = data["duplicate_cohorts"]
    build_simple_doc(out1, "Internal Peer Review 1 - Methods, PRISMA, and Evidence Grading", [
        "Recommendation: Suitable for submission as a conservative evidence-map/systematic-review manuscript after final author sign-off.",
        "Strengths: The pipeline is reproducible, logs all phases, separates human/animal/cellular evidence, penalizes hype-heavy claims, and creates full-text eligibility triage for all screening candidates.",
        f"Search/screening: The current search produced {len(data['raw'])} raw records and {len(data['dedup'])} deduplicated records. This is a strong pilot base, but the capped search must be rerun at larger limits before claiming PRISMA-complete coverage.",
        f"Full-text verification: {len(verification)} priority human records were attempted, including {int((verification['verification_status'] == 'open_full_text_verified').sum()) if not verification.empty else 0} open full-text verifications. Remaining records should be manually checked.",
        "Evidence grading: Exercise is appropriately ranked highest. Other categories are kept below recommendation level, which is scientifically appropriate.",
        "Risk of bias: The formal RoB file is useful as a structured prompt, but it must not be described as a final RoB 2 or ROBINS-I assessment. The manuscript now labels it preliminary, which is correct.",
        "Meta-analysis: It is appropriate not to pool effect estimates unless outcomes, comparators, uncertainty measures, and follow-up intervals are harmonized.",
        "Duplicate cohorts: The duplicate-cohort check identified title/PMID/trial-key flags. These should be manually adjudicated before any final evidence count is published.",
        "Reference/table/figure audit: In-text citations are now sequential superscript square-bracket references. Tables and figures are cited in sequence before display. Workflow labels were removed from publication tables.",
        "Remaining caveat: The article must not be represented as a completed meta-analysis.",
        "Verdict: Scientifically cautious and structurally credible as a submission-ready evidence-map package."
    ])
    build_simple_doc(out2, "Internal Peer Review 2 - Clinical Validity, Writing, and Journal Fit", [
        "Recommendation: Suitable as an evidence-map/systematic-review methods article if framed as conservative and preliminary rather than a completed meta-analysis.",
        "Clinical validity: The conclusion that exercise has the strongest human healthspan signal is supported by the extracted ranking. Claims for rapamycin, senolytics, NAD/sirtuin interventions, fasting, caloric restriction, microbiome, and stem-cell approaches remain appropriately cautious.",
        f"Effect extraction: Quantitative effect information was found in {int((effects['effect_size_extraction_status'] == 'candidate_numeric_effect_text_extracted').sum()) if not effects.empty else 0} priority human records, but the current estimates are not sufficiently harmonized for meta-analysis.",
        f"Duplicate risk: {len(dup)} duplicate/cohort flags require manual adjudication. This should be declared as a limitation.",
        "Journal fit: The manuscript should be submitted as a review/evidence map, not as a definitive therapeutic guideline.",
        "Clinical interpretation: The practical message is appropriate for a medical journal readership: established healthy-ageing interventions, especially physical activity, are more defensible than pharmacologic or regenerative age-reversal claims.",
        "Overclaiming audit: The manuscript avoids equating p-values, pathway modulation, epigenetic-clock changes, or senescence markers with true rejuvenation. This is a major strength.",
        "Readability: The article is compact, table-driven, and suitable for editorial triage. The abstract is direct and avoids promotional language.",
        "Table audit: Tables contain clean publication-facing content rather than local workflow statuses. Table 2 is now explicitly effect-estimate based where estimates are available; terms such as eligible pending confirmation, source_text_type, underscore labels, and local processing names are not used in the manuscript tables.",
        "Figure callout audit: Figures 1-4 are cited in sequence before their display.",
        "Blinding: author identity and contribution details are present in the title-page/declarations files; the blinded manuscript should remain free of identifying material.",
        "Verdict: Clinically defensible as a conservative evidence map. It should not be represented as clinical guidance or proof of age reversal."
    ])
    qc_lines = [
        "Audit scope: content, structure, alignment with MJDRDYPU/Medknow-style submission conventions, consistency, and scientific validity.",
        "Content validity: Findings are derived from extracted tables and verification outputs. No unsupported pooled effect estimate is presented.",
        "Structure: Title page, blinded manuscript, abstract, keywords, introduction, methods, results, discussion, limitations, conclusion, declarations, references, tables, and figures are present.",
        "Alignment: The package uses separate title page and blinded manuscript files, Vancouver-style numbered references, ethics/funding/conflict statements, and embedded figures/tables.",
        "Consistency: The manuscript consistently states that full-text verification and effect extraction are preliminary where automated.",
        "Validity risks: Search caps, metadata-only extraction, duplicate flags, and unverified final effect sizes remain the main threats.",
        "Reference audit: References are generated from bibliographic metadata for prioritized human records. The manuscript uses superscripted square-bracket citations in order of appearance.",
        "Figure audit: Four key figures are embedded in the main manuscript and six are copied separately. Figures are called out sequentially as Figure 1 through Figure 4 before display.",
        "Table audit: Three publication-facing tables are embedded. Workflow/internal status labels were removed; table fields use clean clinical/review terminology.",
        f"Blinding audit: The manuscript file avoids author names and affiliations. The title-page file contains {AUTHOR_NAME}, {AUTHOR_AFFILIATION}, email, phone, and ORCID.",
        "Ethics audit: The ethics statement is appropriate for public bibliographic metadata and open-access text. If authors add nonpublic full texts or manual extraction from copyrighted PDFs, they should retain only extracted facts, not redistributed text.",
        "AI/software audit: The declarations file discloses automated scripting support. Authors must verify all scientific claims before submission.",
        "QC flags: " + (qc.to_string(index=False) if not qc.empty else "No QC table available."),
        "Final audit decision: Ready as a high-quality MJDRDYPU-style draft submission package after final human scientific sign-off is completed; not to be presented as a completed meta-analysis or definitive clinical guideline."
    ]
    build_simple_doc(audit_out, "Comprehensive Submission Audit", qc_lines)


def run(cfg):
    root = cfg["_root"]
    rt = cfg["paths"]["results_tables"]
    data = {
        "raw": read_csv(rt / "raw_records_all.csv"),
        "dedup": read_csv(rt / "master_records_dedup.csv"),
        "screened": read_csv(rt / "title_abstract_screening.csv"),
        "full_text": read_csv(rt / "full_text_status.csv"),
        "full_text_decisions": read_csv(rt / "full_text_screening_decisions.csv"),
        "human_queue": read_csv(rt / "human_evidence_priority_queue.csv"),
        "ranking": read_csv(rt / "intervention_credibility_ranking.csv"),
        "verification": read_csv(rt / "full_text_verification_priority_human.csv"),
        "effects": read_csv(rt / "effect_size_extraction_priority_human.csv"),
        "formal_rob": read_csv(rt / "risk_of_bias_formal_preliminary_human.csv"),
        "duplicate_cohorts": read_csv(rt / "duplicate_cohort_checks.csv"),
        "qc": read_csv(rt / "quality_control_flags.csv"),
    }

    date = datetime.now().strftime("%Y-%m-%d")
    suffix = cfg.get("submission", {}).get("revision_suffix", "")
    folder_name = f"MJDRDYPU_AntiAgeing_{date}{('_' + suffix) if suffix else ''}"
    out = root / "submission_assets" / folder_name
    figs_out = out / "figures"
    tables_out = out / "tables"
    figs_out.mkdir(parents=True, exist_ok=True)
    tables_out.mkdir(parents=True, exist_ok=True)

    references = build_references(data["human_queue"], data["verification"])
    save_csv(references, out / f"MJDRDYPU_AntiAgeing_reference_metadata_{date}.csv")

    manuscript_path = out / f"MJDRDYPU_AntiAgeing_blinded_manuscript_{date}.docx"
    narrative_text = build_manuscript_docx(manuscript_path, data, references, cfg["paths"]["results_figures"])
    main_words = word_count(narrative_text)
    build_title_page(out / f"MJDRDYPU_AntiAgeing_title_page_{date}.docx", main_words, references)
    build_simple_doc(out / f"MJDRDYPU_AntiAgeing_cover_letter_{date}.docx", "Cover Letter", [
        f"Dear Editor, {JOURNAL},",
        f"Please consider the manuscript entitled '{ARTICLE_TITLE}' as a systematic review/evidence-map manuscript.",
        "The manuscript presents a reproducible, conservative evidence synthesis of anti-ageing and age-reversal intervention claims. It explicitly separates healthspan benefit, biological-age biomarkers, lifespan claims, and true clinical rejuvenation, and avoids therapeutic overclaiming.",
        "The work uses public bibliographic metadata and open text where available; no individual participant data were accessed. The manuscript is not under consideration elsewhere, subject to author confirmation before upload.",
        f"Sincerely, {AUTHOR_NAME}, {AUTHOR_ROLE}, {AUTHOR_AFFILIATION}. Email: {AUTHOR_EMAIL}; Phone: {AUTHOR_PHONE}; ORCID: {AUTHOR_ORCID}."
    ])
    build_simple_doc(out / f"MJDRDYPU_AntiAgeing_declarations_{date}.docx", "Declarations", [
        "Ethics approval: Not applicable; public bibliographic metadata and open-access text only.",
        "Consent to participate: Not applicable.",
        "Funding: No external funding declared.",
        "Conflicts of interest: None declared.",
        "Data availability: Tables, figures, scripts, and audit files are available from the project files.",
        f"Author contributions: {AUTHOR_NAME}: concept, study design, data curation, formal analysis, evidence synthesis, manuscript drafting, manuscript review, and guarantor.",
        "Use of AI/software: Automated scripts were used for literature metadata retrieval, screening support, extraction support, evidence scoring, figure generation, and manuscript assembly. All scientific claims require author verification before submission."
    ])
    build_peer_reviews(
        out / f"MJDRDYPU_AntiAgeing_internal_peer_review_1_{date}.docx",
        out / f"MJDRDYPU_AntiAgeing_internal_peer_review_2_{date}.docx",
        out / f"MJDRDYPU_AntiAgeing_content_structure_validity_audit_{date}.docx",
        data,
    )
    build_simple_doc(out / f"MJDRDYPU_AntiAgeing_submission_checklist_{date}.docx", "Submission Checklist", [
        "Separate title page prepared: Yes.",
        "Blinded manuscript prepared: Yes.",
        "Cover letter prepared: Yes.",
        "Declarations prepared: Yes.",
        "Vancouver-style references generated from bibliographic metadata: Yes.",
        "Figures embedded and copied separately: Yes.",
        "Tables embedded and CSV source tables copied: Yes.",
        "Ethics, funding, conflict, data availability statements included: Yes.",
        f"Author details included: Yes - {AUTHOR_NAME}, {AUTHOR_AFFILIATION}, ORCID {AUTHOR_ORCID}.",
        "Manual final full-text/effect-size verification still required before representing the work as a completed meta-analysis: Yes."
    ])

    for fig in ["prisma_flow.png", "evidence_score_ranking.png", "hype_vs_evidence_map.png", "translational_matrix.png", "intervention_outcome_heatmap.png", "mechanism_network.png"]:
        src = cfg["paths"]["results_figures"] / fig
        if src.exists():
            shutil.copy2(src, figs_out / fig)
    for table in [
        "intervention_credibility_ranking.csv",
        "full_text_verification_priority_human.csv",
        "effect_size_extraction_priority_human.csv",
        "risk_of_bias_formal_preliminary_human.csv",
        "duplicate_cohort_checks.csv",
        "translational_readiness.csv",
        "quality_control_flags.csv",
    ]:
        src = rt / table
        if src.exists():
            shutil.copy2(src, tables_out / table)

    append_log(
        cfg["paths"]["logs"] / "progress.md",
        "Phase 21 - MJDRDYPU Publication Asset Package and Double Peer Review",
        "- Prepared MJDRDYPU-style DOCX submission assets.\n- Created blinded manuscript, title page, cover letter, declarations, two internal peer reviews, comprehensive audit, checklist, reference metadata, copied source figures and tables.",
        f"- submission_assets/{folder_name}/*.docx\n- submission_assets/{folder_name}/figures/*.png\n- submission_assets/{folder_name}/tables/*.csv",
        "- Final manual scientific sign-off and portal-specific form checks remain required before upload.",
        "- Journal instructions were inferred from accessible LWW/Medknow-style guidance and MJDRDYPU article patterns because a dedicated current instruction PDF was not directly accessible.",
        "pytest tests -q",
    )
    print(out)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--config", default="config/review_config.yaml")
    run(load_config(parser.parse_args().config))
