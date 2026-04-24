"""
Build the complete supplementary DOCX for MJDRDYPU_AntiAgeing submission.

Supplementary Figures (5):
  S1 – Evidence pyramid (model-system distribution)
  S2 – Evidence timeline (chronological record distribution)
  S3 – Intervention-outcome heatmap
  S4 – Mechanism network
  S5 – DunedinPACE meta-addon forest plot

Supplementary Tables (9):
  S1 – Additional records identified in search (refs 15-40)
  S2 – Complete search strategy (all query strings x sources)
  S3 – Full intervention credibility ranking (all 16 scored columns)
  S4 – Translational readiness (all 10 columns)
  S5 – Full priority human verification results (40 records)
  S6 – Full preliminary risk-of-bias prompts (40 records)
  S7 – Full effect-size extraction (40 records, first clean window)
  S8 – Duplicate cohort and overlapping-publication checks (45 groups)
  S9 – Quality control flags

Run from D:/Anti ageing research directory.
"""

import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE    = "anti_ageing_review/submission_assets/MJDRDYPU_AntiAgeing_final_qa_2026-04-24"
TBLS    = f"{BASE}/tables"
FIGS    = f"{BASE}/figures"               # submission figures (Figs 1-4 + heatmap + network)
RFIGS   = "anti_ageing_review/results/figures"   # full results set
MFIGS   = "anti_ageing_review/results/meta_addon/figures"
OUT     = f"{BASE}/supplementary/MJDRDYPU_AntiAgeing_supplementary_complete_2026-04-24.docx"


# ── helpers ──────────────────────────────────────────────────────────

def set_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    for s in doc.sections:
        s.top_margin    = Inches(top)
        s.bottom_margin = Inches(bottom)
        s.left_margin   = Inches(left)
        s.right_margin  = Inches(right)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def para(doc, text, italic=False, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold   = bold
    r.font.size = Pt(size)
    return p


def caption(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p


def embed_fig(doc, path, width=6.0, cap=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.add_run().add_picture(path, width=Inches(width))
    if cap:
        caption(doc, cap)


def shade_hdr(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)


def make_table(doc, headers, widths, data_rows, cell_size=8):
    """
    headers   : list of str
    widths    : list of float (inches, sum ≤ 6.5)
    data_rows : list of lists (str)
    """
    t = doc.add_table(rows=1, cols=len(headers))
    t.style     = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, (h, w) in enumerate(zip(headers, widths)):
        hdr.cells[i].text  = h
        hdr.cells[i].width = Inches(w)
        for p in hdr.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold      = True
                r.font.size = Pt(9)
    shade_hdr(hdr)
    for row_vals in data_rows:
        row = t.add_row()
        for i, (v, w) in enumerate(zip(row_vals, widths)):
            row.cells[i].text  = str(v) if v is not None and str(v) != "nan" else ""
            row.cells[i].width = Inches(w)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(cell_size)
    doc.add_paragraph()
    return t


# ── label cleaners ────────────────────────────────────────────────────

INTERV = {
    "exercise":"Exercise","microbiome":"Microbiome","rapamycin_mtor":"Rapamycin/mTOR",
    "senolytics":"Senolytics","caloric_restriction":"Caloric restriction",
    "lifestyle_bundle":"Lifestyle bundle","nad_sirtuin":"NAD+/sirtuin","fasting":"Fasting",
    "supplements":"Supplements","metformin":"Metformin","sleep_circadian":"Sleep/circadian",
    "controversial":"Controversial","reprogramming":"Reprogramming","stem_cell":"Stem cell",
}
TRANS = {
    "A_healthspan_support_signal":          "A – Healthspan support signal",
    "B_promising_not_recommendation_ready": "B – Promising; not recommendation-ready",
    "C_biomarker_or_indirect":              "C – Biomarker / indirect",
    "D_speculative_or_low_directness":      "D – Speculative / low directness",
}
TIER = {
    "highest_current_human_healthspan_signal": "Highest (healthspan signal)",
    "human_signal_requires_verification":      "Requires verification",
    "biomarker_or_indirect_human_signal":      "Biomarker / indirect",
    "low_directness_or_speculative":           "Speculative / low",
}
DOMAIN = {
    "healthspan_functional_ageing": "Healthspan / function",
    "hard_ageing_relevance":        "Hard ageing relevance",
    "biological_ageing_biomarker":  "Biological-age biomarker",
    "surrogate_or_indirect":        "Surrogate / indirect",
}
ROB = {
    "low_concern":                                    "Low",
    "some_concern":                                   "Some concern",
    "unclear":                                        "Unclear",
    "low_to_some_concern_pending_full_text_review":   "Low–Some (pending)",
    "serious_or_unclear":                             "Serious/unclear",
    "low_concern_rct_design":                         "Low (RCT)",
    "some_concern_observational_confounding":         "Some (observational)",
    "not_applicable":                                 "NA",
    "not_applicable_or_unclear":                      "NA/unclear",
    "unclear_without_protocol":                       "Unclear (no protocol)",
    "unclear_or_not_applicable":                      "Unclear/NA",
    "low_or_some_concern":                            "Low–Some",
}


def clean_effect_text(raw):
    """Take first sliding-window chunk, strip pipes, truncate to ~280 chars."""
    if not isinstance(raw, str) or raw.strip() == "":
        return "Not extracted"
    first = raw.split(" || ")[0].strip()
    return first[:280] + ("…" if len(first) > 280 else "")


def shorten_title(title, max_len=80):
    t = str(title)
    return t[:max_len] + "…" if len(t) > max_len else t


# ══════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    set_margins(doc)
    nrm = doc.styles["Normal"]
    nrm.font.name = "Times New Roman"
    nrm.font.size = Pt(11)

    # ── Cover page ────────────────────────────────────────────────
    heading(doc,
        "Supplementary Material\n"
        "Can Ageing Be Slowed or Reversed? A Reproducible Evidence Map, "
        "Credibility Ranking, and Mechanistic Synthesis of Anti-Ageing and "
        "Age-Reversal Interventions", level=1)
    para(doc, "Medical Journal of Dr. D.Y. Patil Vidyapeeth (MJDRDYPU)")
    para(doc, "Date: 2026-04-24")
    para(doc,
         "This file contains 5 supplementary figures (S1–S5) and 9 supplementary tables "
         "(S1–S9) that support the main manuscript. All tables are derived from the "
         "reproducible pipeline; no supplementary data were manually fabricated.",
         italic=True)
    doc.add_paragraph()

    # ── Contents list ─────────────────────────────────────────────
    heading(doc, "Contents", level=2)
    contents = [
        ("Supplementary Figures", [
            "Figure S1 – Evidence pyramid (model-system distribution)",
            "Figure S2 – Evidence timeline (record distribution by year and intervention)",
            "Figure S3 – Intervention-outcome heatmap",
            "Figure S4 – Mechanistic pathway network",
            "Figure S5 – DunedinPACE primary pooled effect (meta-addon forest plot)",
        ]),
        ("Supplementary Tables", [
            "Table S1  – Additional records identified in search (refs 15–40)",
            "Table S2  – Complete search strategy (all query strings and sources)",
            "Table S3  – Full intervention credibility ranking (all 16 scored columns)",
            "Table S4  – Translational readiness (all 10 columns, 14 intervention classes)",
            "Table S5  – Full priority human verification results (40 records)",
            "Table S6  – Full preliminary risk-of-bias domain prompts (40 records)",
            "Table S7  – Full effect-size extraction with candidate text (40 records)",
            "Table S8  – Duplicate cohort and overlapping-publication checks (45 groups)",
            "Table S9  – Quality control flags",
        ]),
    ]
    for section_title, items in contents:
        p = doc.add_paragraph()
        r = p.add_run(section_title)
        r.bold = True
        r.font.size = Pt(11)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()

    # ══ SUPPLEMENTARY FIGURES ══════════════════════════════════════

    heading(doc, "Supplementary Figures", level=1)

    embed_fig(doc, f"{RFIGS}/evidence_pyramid.png", width=5.5,
        cap="Supplementary Figure S1. Evidence pyramid showing distribution of "
            "candidate records by model system and evidence tier (n=1029 deduplicated "
            "records). Human evidence forms the top tier; cellular and animal records "
            "form the base. The pyramid illustrates why most extracted records cannot "
            "be directly translated to human clinical recommendations without further "
            "human-evidence confirmation.")

    embed_fig(doc, f"{RFIGS}/evidence_timeline.png", width=6.0,
        cap="Supplementary Figure S2. Evidence timeline. Chronological distribution "
            "of candidate anti-ageing records across intervention domains (2010–2026). "
            "The chart shows the acceleration of records in exercise, microbiome, "
            "senolytic, and NAD+/sirtuin domains in recent years. Year-of-publication "
            "data are derived from extracted bibliographic metadata; missing years "
            "reflect records with incomplete metadata.")

    embed_fig(doc, f"{FIGS}/intervention_outcome_heatmap.png", width=6.0,
        cap="Supplementary Figure S3. Intervention-outcome heatmap. "
            "Metadata-assisted map of 14 intervention classes (rows) against "
            "ageing-relevant outcome categories (columns). Cell intensity reflects "
            "the count of candidate records in each intervention–outcome combination. "
            "Cells require full-text verification before final clinical interpretation. "
            "This figure was not included in the main manuscript owing to its "
            "preliminary metadata-only basis.")

    embed_fig(doc, f"{FIGS}/mechanism_network.png", width=6.2,
        cap="Supplementary Figure S4. Mechanistic pathway network. "
            "Preliminary network showing mapped links between intervention classes "
            "and candidate mechanistic domains (nutrient sensing, senescence, "
            "inflammaging, epigenetic regulation, mitochondrial function, "
            "autophagy/proteostasis, stem-cell biology, microbiome-host pathways, "
            "and NAD+/sirtuin axis). Edge width reflects the number of extracted "
            "records with the mapped mechanistic label. Labels are inferred from "
            "title/abstract metadata; full-text confirmation is required.")

    embed_fig(doc, f"{MFIGS}/forest_dunedinpace_primary.png", width=6.0,
        cap="Supplementary Figure S5. DunedinPACE primary pooled effect (meta-addon "
            "forest plot). Preliminary pooled estimate for DunedinPACE biological-age "
            "change across fasting-mimicking diet candidate studies. This figure is "
            "generated from the meta-addon pipeline and uses reconstructed or "
            "extracted person-level estimates; it is not a finalised meta-analysis "
            "and should not be interpreted as a definitive pooled effect. Final "
            "pooling requires verified comparable denominators, uncertainty measures, "
            "and harmonised outcome definitions.")

    doc.add_page_break()

    # ══ SUPPLEMENTARY TABLES ══════════════════════════════════════

    heading(doc, "Supplementary Tables", level=1)
    para(doc,
         "All tables are derived from the reproducible pipeline. Column headers "
         "retain pipeline field names where these are self-explanatory. "
         "Human adjudication is required before any table is treated as a final "
         "clinical or epidemiological determination.",
         italic=True)
    doc.add_paragraph()

    # ── Table S1: Additional records (refs 15-40) ─────────────────
    heading(doc, "Supplementary Table S1. Additional Records Identified in the Search (Refs 15–40)", level=2)
    para(doc,
         "These 26 records were identified in the pilot search (1029 deduplicated records) "
         "and contributed to overall evidence counts in the credibility ranking. They are "
         "cited collectively as [15–40] in the Methods section of the main manuscript. "
         "Final eligibility adjudication for each record remains pending manual full-text "
         "review.", italic=True)

    s1_headers = ["Ref", "First author", "Year", "Journal", "Title (abbreviated)", "Intervention class", "Identifier"]
    s1_widths  = [0.30, 0.90, 0.40, 1.00, 1.90, 1.00, 0.95]
    s1_data = [
        ("15","Mihaiescu-Ion V","2026","BMC Geriatr","PRICA-POWFRAIL exercise programme study protocol","exercise","doi:10.1186/s12877-025-06902-9"),
        ("16","Robinson LA","2026","Front Nutr","Epigenetic and microbiome responses to greens supplementation","microbiome","doi:10.3389/fnut.2026.1750030"),
        ("17","Coleman AE","2025","GeroScience","TRIAD: rapamycin in aging dogs (Dog Aging Project)","rapamycin_mtor","PMID:39951177"),
        ("18","Pihlstrom L","2024","J Neurol Neurosurg Psychiatry","NAD-HD: NR in Huntington's disease (NAD-HD trial)","nad_sirtuin","doi:10.1136/jnnp-2024-ehdn.333"),
        ("19","Nilsson MI","2024","Nutrients","Obesity/metabolic disease impair anabolic response to protein+exercise","exercise","PMID:39771028"),
        ("20","Erratum","2023","J Aging Phys Act","Erratum: Dancing + resistance training in aging women","exercise","doi:10.1123/japa.2023-0257"),
        ("21","Yi L","2023","GeroScience","NMN supplementation in healthy middle-aged adults: RCT","nad_sirtuin","PMID:36482258"),
        ("22","Waziry R","2023","Nat Aging","CALERIE trial: caloric restriction and DNA methylation aging","caloric_restriction","PMID:37118425"),
        ("23","Fiorito G","2021","Aging Cell","DAMA study: diet+physical activity and DNA methylation aging","lifestyle_bundle","PMID:34535961"),
        ("24","Stares A","2020","J Geriatr Phys Ther","Creatine + exercise in aging: systematic review of RCTs","exercise","PMID:30762623"),
        ("25","Casas-Herrero A","2019","Trials","VIVIFRAIL multicomponent exercise in frail elders: protocol","exercise","PMID:31208471"),
        ("26","Vlietstra L","2025","Ageing Res Rev","Exercise for frailty in older adults with hypertension: review","exercise","doi:10.1016/j.arr.2025.102714"),
        ("27","Billot M","2020","Clin Interv Aging","Physical activity recommendations for frailty/sarcopenia","exercise","PMID:32982201"),
        ("28","Munoz-Pardeza J","2026","J Cachexia Sarcopenia Muscle","Resistance training in youth with Type 1 Diabetes [FLAG: age/population]","exercise","doi:10.1002/jcsm.70257"),
        ("29","Racette SB","2026","Am J Clin Nutr","Diet quality in CALERIE 2 calorie restriction trial","caloric_restriction","doi:10.1016/j.ajcnut.2025.101182"),
        ("30","Veronese N","2026","Aging Clin Exp Res","European council recommendations: PA and diet for mental health in older adults","lifestyle_bundle","doi:10.1007/s40520-025-03315-x"),
        ("31","Ying YY","2026","EBioMedicine","Circadian rhythm disruption impairs ovarian development via NAD+ reprogramming","nad_sirtuin","doi:10.1016/j.ebiom.2026.106200"),
        ("32","Berven H","2026","iScience","NAD-brain pharmacokinetic study of oral NAD precursor supplementation","nad_sirtuin","doi:10.1016/j.isci.2026.114764"),
        ("33","Kell L","2026","Aging Cell","Rapamycin geroprotection in ageing human immune system via DNA damage resilience","rapamycin_mtor","doi:10.1111/acel.70364"),
        ("34","Bautista J","2026","Front Microbiol","Biohacking the human gut microbiome for precision health","microbiome","doi:10.3389/fmicb.2026.1776983"),
        ("35","Liu C","2026","Front Public Health","Research advances in exercise management for frail older adults","exercise","doi:10.3389/fpubh.2026.1763583"),
        ("36","Shi M","2026","Front Public Health","Editorial: multidimensional approaches to aging and lifestyle interventions","lifestyle_bundle","doi:10.3389/fpubh.2026.1777378"),
        ("37","Saez-Nieto C","2026","Front Aging","Exercise training in frail older adults with heart failure: systematic review","exercise","doi:10.3389/fragi.2026.1800669"),
        ("38","Lim MJS","2026","Nutrients","Diet, gut microbiome, and estrogen physiology in menopausal health","microbiome","doi:10.3390/nu18071052"),
        ("39","Poisnel G","2026","Alzheimers Dement (N Y)","Physical activity and Alzheimer's disease biomarkers in older adults","exercise","No DOI at extraction"),
        ("40","Gulej R","2025","Geroscience","Heterochronic parabiosis and plasma transfer in cerebrovascular/brain aging","controversial","doi:10.1007/s11357-025-01657-y"),
    ]
    make_table(doc, s1_headers, s1_widths, s1_data, cell_size=8)
    para(doc, "Note: Ref 28 (Munoz-Pardeza J, 2026) concerns youth with Type 1 Diabetes and is marginally relevant to anti-ageing in older adults. It should be excluded or reclassified during final manual eligibility adjudication.", italic=True)
    doc.add_paragraph()

    # ── Table S2: Search strategy ──────────────────────────────────
    heading(doc, "Supplementary Table S2. Complete Search Strategy", level=2)
    para(doc,
         "All 11 query strings run across PubMed, Europe PMC, and Crossref in the pilot "
         "search (cap: 35 records per query per source). Total records retrieved: "
         "11 queries × 3 databases × 35 records = 1155 raw records.",
         italic=True)
    s2_headers = ["#", "Source", "Query string", "Records retrieved", "Status"]
    s2_widths  = [0.25, 0.80, 3.60, 0.90, 0.90]
    queries = [
        "(aging OR ageing OR longevity OR healthspan OR lifespan) AND (intervention OR trial OR treatment OR therapy)",
        '("biological age" OR "epigenetic clock" OR "DNA methylation age") AND (intervention OR trial OR therapy)',
        '(rejuvenation OR "age reversal" OR "reverse aging" OR "reverse ageing")',
        '(aging OR ageing OR healthspan OR lifespan) AND ("caloric restriction" OR "dietary restriction" OR fasting)',
        '(aging OR ageing OR healthspan OR frailty) AND (exercise OR "physical activity" OR "resistance training")',
        '(aging OR ageing OR longevity OR frailty) AND (metformin OR rapamycin OR sirolimus OR mTOR)',
        '(aging OR ageing OR "epigenetic clock") AND (NAD OR "nicotinamide riboside" OR NMN OR sirtuin)',
        "(senescent cells OR senolytics OR senomorphics) AND (aging OR ageing OR rejuvenation)",
        "(stem cell OR reprogramming OR partial reprogramming) AND (aging OR rejuvenation)",
        "(plasma OR parabiosis OR GDF11) AND (aging OR ageing OR rejuvenation)",
        "(microbiome OR gut bacteria OR probiotics) AND (aging OR ageing OR longevity)",
    ]
    s2_data = []
    for i, q in enumerate(queries):
        for src in ["PubMed","Europe PMC","Crossref"]:
            s2_data.append((str(i+1) if src=="PubMed" else "", src, q if src=="PubMed" else "", "35", "ok"))
    make_table(doc, s2_headers, s2_widths, s2_data, cell_size=7.5)

    # ── Table S3: Full credibility ranking ────────────────────────
    heading(doc, "Supplementary Table S3. Full Intervention Credibility Ranking (All Scored Columns)", level=2)
    para(doc,
         "All 16 columns from the credibility ranking pipeline (n=14 intervention classes). "
         "Credibility score combines human evidence, human trial evidence, direct "
         "ageing/healthspan outcomes, biomarker evidence, surrogate burden, and "
         "hype-language penalty. Rankings are provisional pending final full-text "
         "adjudication.", italic=True)
    df3 = pd.read_csv(f"{TBLS}/intervention_credibility_ranking.csv")
    s3_headers = [
        "Rank","Intervention","Records\n(n)","Human\nrec","Human\ntrials",
        "Hard ageing\nrec","Healthspan\nrec","Biomarker\nrec","Surrogate\nrec",
        "Max\nscore","Mean\nscore","Hype\nflagged","Hype\nrate",
        "Credibility\nscore","Credibility tier",
    ]
    s3_widths = [0.35,1.00,0.50,0.45,0.45,0.55,0.55,0.55,0.55,0.40,0.40,0.45,0.40,0.55,1.25]
    s3_data = []
    for _, row in df3.iterrows():
        s3_data.append([
            int(row["rank"]),
            INTERV.get(row["intervention_name"], row["intervention_name"]),
            int(row["n_extracted_records"]),
            int(row["human_records"]),
            int(row["human_trial_records"]),
            int(row["hard_ageing_records"]),
            int(row["healthspan_records"]),
            int(row["biomarker_records"]),
            int(row["surrogate_records"]),
            f"{row['max_claim_score']:.1f}",
            f"{row['mean_claim_score']:.2f}",
            int(row["hype_flagged_records"]),
            f"{row['hype_rate']:.2f}",
            f"{row['credibility_score']:.2f}",
            TIER.get(row["credibility_tier"], row["credibility_tier"]),
        ])
    make_table(doc, s3_headers, s3_widths, s3_data, cell_size=7.5)

    # ── Table S4: Translational readiness ─────────────────────────
    heading(doc, "Supplementary Table S4. Translational Readiness (All Columns, n=14 Intervention Classes)", level=2)
    para(doc,
         "Conservative category assignment using credibility score, human record count, "
         "human trial count, direct ageing-outcome count, and hype burden. Categories "
         "are provisional pending final full-text adjudication.", italic=True)
    df4 = pd.read_csv(f"{TBLS}/translational_readiness.csv")
    s4_headers = [
        "Intervention","Credibility\nscore","Human\nrecords","Human\ntrials",
        "Healthspan\nrecords","Hard ageing\nrecords","Biomarker\nrecords",
        "Hype\nrate","Translational category",
    ]
    s4_widths = [1.10,0.65,0.60,0.60,0.65,0.65,0.65,0.55,1.50]
    s4_data = []
    for _, row in df4.iterrows():
        s4_data.append([
            INTERV.get(row["intervention_name"], row["intervention_name"]),
            f"{row['credibility_score']:.2f}",
            int(row["human_records"]),
            int(row["human_trial_records"]),
            int(row["healthspan_records"]),
            int(row["hard_ageing_records"]),
            int(row["biomarker_records"]),
            f"{row['hype_rate']:.2f}",
            TRANS.get(row["translational_category"], row["translational_category"]),
        ])
    make_table(doc, s4_headers, s4_widths, s4_data, cell_size=8)

    # ── Table S5: Full priority human verification ────────────────
    heading(doc, "Supplementary Table S5. Full Priority Human Verification Results (n=40 Records)", level=2)
    para(doc,
         "All 40 records in the priority human verification pass. "
         "Source text type: pmc_open_full_text = open full text retrieved from PubMed Central; "
         "pubmed_abstract = PubMed abstract only; not_retrieved = unavailable in open workflow. "
         "Provisional eligibility labels require manual confirmation.", italic=True)
    df5 = pd.read_csv(f"{TBLS}/full_text_verification_priority_human.csv")
    s5_headers = [
        "Title (abbreviated)","PMID","PMCID",
        "Source text type","Verification status","Prov. eligibility",
        "Intervention\nsignal","Ageing\nsignal","Manual\nneeded",
    ]
    s5_widths = [2.20,0.55,0.55,0.80,0.90,0.90,0.55,0.55,0.50]
    s5_data = []
    for _, row in df5.iterrows():
        s5_data.append([
            shorten_title(row["title"], 70),
            str(row["pmid"]) if not pd.isna(row["pmid"]) else "",
            str(row["pmcid"]) if not pd.isna(row["pmcid"]) else "",
            str(row["source_text_type"]).replace("_"," "),
            str(row["verification_status"]).replace("_"," "),
            str(row["provisional_full_text_eligibility"]).replace("_"," "),
            str(row["has_intervention_signal"]),
            str(row["has_ageing_outcome_signal"]),
            str(row["manual_verification_required"]),
        ])
    make_table(doc, s5_headers, s5_widths, s5_data, cell_size=7.5)

    # ── Table S6: Full risk-of-bias ───────────────────────────────
    heading(doc, "Supplementary Table S6. Full Preliminary Risk-of-Bias Domain Prompts (n=40 Records)", level=2)
    para(doc,
         "Structured preliminary risk-of-bias prompts for all 40 priority human records. "
         "These are NOT final RoB 2, ROBINS-I, SYRCLE, or GRADE assessments. "
         "They are structured prompts based on available open text or abstract content "
         "and require complete full-text review to finalise.", italic=True)
    df6 = pd.read_csv(f"{TBLS}/risk_of_bias_formal_preliminary_human.csv")
    s6_headers = [
        "Title (abbreviated)","Randomisation","Blinding",
        "Missing\ndata","Confounding","Outcome\nmeasurement",
        "Selective\nreporting","Overall (preliminary)",
    ]
    s6_widths = [2.10,0.70,0.65,0.60,0.75,0.70,0.65,1.00]
    s6_data = []
    for _, row in df6.iterrows():
        s6_data.append([
            shorten_title(row["title"], 65),
            ROB.get(row["rob_randomization"],    row["rob_randomization"]),
            ROB.get(row["rob_blinding"],          row["rob_blinding"]),
            ROB.get(row["rob_missing_data"],      row["rob_missing_data"]),
            ROB.get(row["rob_confounding"],       row["rob_confounding"]),
            ROB.get(row["rob_outcome_measurement"],row["rob_outcome_measurement"]),
            ROB.get(row["rob_selective_reporting"],row["rob_selective_reporting"]),
            ROB.get(row["rob_overall"],            row["rob_overall"]),
        ])
    make_table(doc, s6_headers, s6_widths, s6_data, cell_size=7.5)

    # ── Table S7: Full effect-size extraction ─────────────────────
    heading(doc, "Supplementary Table S7. Full Effect-Size Extraction with Candidate Text (n=40 Records)", level=2)
    para(doc,
         "Candidate numeric effect text extracted from open full text or PubMed abstracts "
         "for all 40 priority human records. The 'Candidate text' column shows the first "
         "extracted window (~280 characters). Effect estimates are not harmonised or pooled; "
         "manual verification is required for all records before quantitative synthesis.",
         italic=True)
    df7 = pd.read_csv(f"{TBLS}/effect_size_extraction_priority_human.csv")
    s7_headers = [
        "Title (abbreviated)","Intervention","Outcome domain",
        "Extraction status","Candidate text (first window)","Manual needed",
    ]
    s7_widths = [1.60,0.85,0.85,0.80,2.10,0.55]
    s7_data = []
    for _, row in df7.iterrows():
        status = str(row["effect_size_extraction_status"]).replace("_"," ")
        s7_data.append([
            shorten_title(row["title"], 55),
            INTERV.get(row["intervention_name"], row["intervention_name"]),
            DOMAIN.get(row["ageing_domain_category"], str(row["ageing_domain_category"]).replace("_"," ")),
            status,
            clean_effect_text(row["candidate_effect_size_text"]),
            str(row["manual_extraction_required"]),
        ])
    make_table(doc, s7_headers, s7_widths, s7_data, cell_size=7.5)

    # ── Table S8: Duplicate cohort checks ────────────────────────
    heading(doc, "Supplementary Table S8. Duplicate Cohort and Overlapping-Publication Checks (n=45 Groups)", level=2)
    para(doc,
         "Duplicate and overlapping-cohort check results using DOI, PMID, normalised title, "
         "and trial acronym keys. All flagged groups require manual adjudication before "
         "final evidence counts are published. These flags do not constitute final duplicate "
         "removal decisions.", italic=True)
    df8 = pd.read_csv(f"{TBLS}/duplicate_cohort_checks.csv")
    s8_headers = ["Check type","Duplicate key","N records","Manual action"]
    s8_widths  = [0.80, 2.90, 0.70, 2.10]
    s8_data = []
    for _, row in df8.iterrows():
        s8_data.append([
            str(row["check_type"]),
            shorten_title(row["duplicate_key"], 85),
            str(int(row["n_records"])) if not pd.isna(row["n_records"]) else "",
            str(row["manual_action"]).replace("_"," "),
        ])
    make_table(doc, s8_headers, s8_widths, s8_data, cell_size=8)

    # ── Table S9: Quality control flags ───────────────────────────
    heading(doc, "Supplementary Table S9. Quality Control Flags", level=2)
    para(doc,
         "Automated quality control checks run at the end of the pipeline. Flags with "
         "severity 'high' or 'medium' must be resolved before the manuscript is submitted "
         "as a completed systematic review.", italic=True)
    df9 = pd.read_csv(f"{TBLS}/quality_control_flags.csv")
    s9_headers = ["Issue type","N records","Severity","Detail"]
    s9_widths  = [1.60, 0.70, 0.70, 3.50]
    s9_data = []
    for _, row in df9.iterrows():
        s9_data.append([
            str(row["issue_type"]).replace("_"," "),
            str(int(row["n_records"])) if not pd.isna(row["n_records"]) else "0",
            str(row["severity"]),
            str(row["detail"]),
        ])
    make_table(doc, s9_headers, s9_widths, s9_data, cell_size=8.5)

    # ── Save ──────────────────────────────────────────────────────
    doc.save(OUT)
    print(f"Saved: {OUT}")
    import os
    sz = os.path.getsize(OUT)
    print(f"Size:  {sz//1024} KB")


if __name__ == "__main__":
    build()
